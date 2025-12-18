
import json
import os
import sys
import threading
import time
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from typing import Dict, List, Optional, Any
import zipfile
import subprocess
import re
from tempfile import NamedTemporaryFile
from docx.oxml import parse_xml
import traceback

_FILE_LOCK = threading.RLock()
_OUTPUT_DIR_LOCK = threading.RLock()

def get_app_path():
    """Lấy đường dẫn chứa file .exe hoặc script"""
    if getattr(sys, 'frozen', False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.abspath(__file__))
def find_pandoc_executable():
    """
    Tìm pandoc.exe theo thứ tự ưu tiên:
    1. Thư mục 'pandoc' cạnh tool (cho bản build)
    2. PATH hệ thống (cho môi trường dev)
    """
    app_path = get_app_path()
    
    # 1. Tìm trong thư mục cục bộ 'pandoc' (ưu tiên cao nhất)
    local_pandoc = os.path.join(app_path, 'pandoc', 'pandoc.exe')
    if os.path.isfile(local_pandoc):
        # print(f"✅ Sử dụng Pandoc cục bộ: {local_pandoc}")
        return local_pandoc
    
    # 2. Fallback: Tìm trong PATH hệ thống (cho dev)
    import shutil
    system_pandoc = shutil.which('pandoc')
    if system_pandoc:
        # print(f"⚠️ Sử dụng Pandoc hệ thống: {system_pandoc}")
        return system_pandoc
    
    # 3. Không tìm thấy
    print("❌ KHÔNG TÌM THẤY PANDOC!")
    return None

def latex_to_omml_via_pandoc(latex_math_dollar):
    """Chuyển đổi LaTeX sang OMML qua Pandoc"""
    pandoc_exe = find_pandoc_executable()
    
    if not pandoc_exe:
        print("❌ Pandoc không khả dụng, bỏ qua equation")
        return None
    
    try:
        # Chuẩn hóa input (loại bỏ ký tự lạ)
        latex_clean = latex_math_dollar.strip()
        
        # Tạo file tạm với encoding UTF-8 BOM để tránh lỗi
        with NamedTemporaryFile(mode='w', suffix=".docx", delete=False, encoding='utf-8') as temp_docx:
            temp_path = temp_docx.name
        
        # Chạy Pandoc với error handling tốt hơn
        result = subprocess.run(
            [pandoc_exe, '--from=latex', '--to=docx', '-o', temp_path],
            input=latex_clean,
            text=True,
            encoding='utf-8',
            capture_output=True,
            timeout=10,  # Timeout 10s để tránh treo
            creationflags=subprocess.CREATE_NO_WINDOW if hasattr(subprocess, 'CREATE_NO_WINDOW') else 0
        )
 
        if result.returncode != 0:
            error_msg = result.stderr.strip() if result.stderr else "Unknown error"
            print(f"⚠️ Pandoc error (code {result.returncode}): {error_msg}")
            
            # Kiểm tra lỗi phổ biến
            if "not found" in error_msg.lower() or "cannot find" in error_msg.lower():
                print("   → Thiếu DLL dependencies. Kiểm tra lại folder pandoc/")
            elif "syntax" in error_msg.lower():
                print(f"   → LaTeX syntax error: {latex_clean[:50]}...")
            
            return None
        
        # Kiểm tra file output có tồn tại không
        if not os.path.exists(temp_path) or os.path.getsize(temp_path) == 0:
            print(f"⚠️ Pandoc không tạo file output hợp lệ")
            return None
           
        # Đọc XML từ DOCX
        with zipfile.ZipFile(temp_path, 'r') as z:
            xml_content = z.read('word/document.xml').decode('utf-8')
        
        # Dọn dẹp file tạm
        try:
            os.remove(temp_path)
        except:
            pass
       
        # Tìm equation XML
        match = re.search(r'(<m:oMath[^>]*>.*?</m:oMath>)', xml_content, re.DOTALL)
        
        if not match:
            print(f"⚠️ Không tìm thấy equation trong output: {latex_clean[:30]}...")
            return None
            
        return match.group(1)
   
    except subprocess.TimeoutExpired:
        print(f"⚠️ Pandoc timeout (>10s)")
        return None
    except Exception as e:
        print(f"❌ Lỗi latex_to_omml: {type(e).__name__}: {e}")
        import traceback
        traceback.print_exc()
        return None



def process_text_with_latex(text, paragraph, bold=False):
    """
    Xử lý text có công thức LaTeX
    VERSION ỔN ĐỊNH - Copy từ test_res.py (KHÔNG có repair_broken_latex)
    """
    if not text:
        return
    
    # Làm sạch HTML tags
    text = text.replace("<br>", "\n").replace("<br/>", "\n") \
               .replace("<Br>", "\n").replace("<Br/>", "\n")
    text = re.sub(r'</?(div|p|u|span|font|i|b)\b[^>]*>', '', text)
    text = text.replace("&nbsp;", "").replace("&lt;", "").replace("&gt;", "")
    
    # Tách text và LaTeX
    pattern = r'(\$[^$]+\$|\\\[.*?\\\])'
    parts = re.split(pattern, text)
    
    for part in parts:
        if not part:
            continue
        
        # Phần LaTeX
        if part.startswith('$') or part.startswith('\\['):
            try:
                latex_expr = clean_latex_math(part)
                insert_equation_into_paragraph(latex_expr, paragraph)
            except Exception as e:
                # Fallback: thêm text thuần
                run = paragraph.add_run(part)
                if bold:
                    run.bold = True
        # Phần text thường
        else:
            cleaned_part = re.sub(r'^\s*/', '', part)
            run = paragraph.add_run(cleaned_part)
            if bold:
                run.bold = True


def insert_equation_into_paragraph(latex_math_dollar, paragraph):
    """Chèn công thức toán học vào paragraph"""
    omml_str = latex_to_omml_via_pandoc(latex_math_dollar)
    
    if not omml_str:
        # Fallback: Thêm text thuần nếu không convert được
        paragraph.add_run(f" [{latex_math_dollar}] ")
        return
    
    # Thêm namespace nếu thiếu
    if 'xmlns:m=' not in omml_str:
        omml_str = re.sub(
            r'<m:oMath',
            r'<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"',
            omml_str,
            count=1
        )
    
    try:
        omml_element = parse_xml(omml_str)
        run = paragraph.add_run()
        run._r.append(omml_element)
    except Exception as e:
        print(f"Lỗi chèn equation: {e}")
        paragraph.add_run(f" [{latex_math_dollar}] ")


def clean_latex_math(latex_raw):
    latex_raw = re.sub(r'\\/', '', latex_raw)
    latex_raw = re.sub(r'\\operatorname\s*{\s*([^}]*)\s*}',
                       lambda m: m.group(1).replace(' ', ''), latex_raw)
    latex_raw = re.sub(r'\\root\s*(\d+)\s*{([^}]*)}', r'\\sqrt[\1]{\2}', latex_raw)
    latex_raw = re.sub(r'\\root\s*{(\d+)}\s*\\of\s*{([^}]*)}', r'\\sqrt[\1]{\2}', latex_raw)
    latex_raw = re.sub(r'\\root\s*(\d+)\s*\\sqrt\s*{([^}]*)}', r'\\sqrt[\1]{\2}', latex_raw)
    latex_raw = re.sub(r'([a-zA-Z])\s*\\frac\s*{([^}]+)}\s*{([^}]+)}',
                       r'\1^{\\frac{\2}{\3}}', latex_raw)
    latex_raw = re.sub(r'\\sp\s*{([^}]*)}', r'^{\1}', latex_raw)
    latex_raw = re.sub(r'{\\bf\s*([^}]*)}', r'\1', latex_raw)
    latex_raw = re.sub(r'\\\s*log', r'\\log', latex_raw)
    latex_raw = re.sub(r'\\bigskip', '', latex_raw)
    latex_raw = re.sub(r'\\nonumber', '', latex_raw)
    latex_raw = latex_raw.replace(r'\?', '?')
    latex_raw = re.sub(r'\\cdot\s*(?=\w)', r'\\cdot ', latex_raw)
    latex_raw = latex_raw.replace(r'\dotstan', r'\cdot \tan')
    latex_raw = re.sub(r'(?<!\\)(\bln\b|\blog\b|\bsin\b|\bcos\b|\btan\b|\blog_{?\d*}?)',
                       r'\\\1', latex_raw)
    latex_raw = re.sub(r'(\\Leftrightarrow|\\Rightarrow|\\rightarrow)(?=\w)', r'\1 ', latex_raw)
    latex_raw = latex_raw.replace(r'\\n', r'\n')
    
    latex_raw = latex_raw.strip()
    # ✅ KHÁC BIỆT: Version cũ KHÔNG replace \n và \r
    # latex_raw = latex_raw.replace('\n', ' ').replace('\r', '')  # ← XÓA DÒNG NÀY
    
    if not (latex_raw.startswith('$') and latex_raw.endswith('$')):
        latex_raw = f"${latex_raw}$"
    
    return latex_raw

def ensure_output_folder_for_batch(batch_name):
    """Tạo folder riêng cho batch"""
    base_path = get_app_path()
    output_base = os.path.join(base_path, "output")
    batch_folder = os.path.join(output_base, batch_name)
    
    with _OUTPUT_DIR_LOCK:
        os.makedirs(output_base, exist_ok=True)
        os.makedirs(batch_folder, exist_ok=True)
    
    return batch_folder

def save_document_securely(doc, batch_name, file_name):
    """Lưu file DOCX với thread-safety"""
    batch_folder = ensure_output_folder_for_batch(batch_name)
    if not batch_folder:
        return None

    output_path = os.path.join(batch_folder, f"{file_name}.docx")
    
    with _FILE_LOCK:
        max_retries = 3
        for retry_count in range(max_retries):
            try:
                doc.save(output_path)
                if os.path.exists(output_path):
                    file_size = os.path.getsize(output_path)
                    print(f"✅ Đã lưu file: {output_path} ({file_size} bytes)")
                    return output_path
            except Exception as e:
                print(f"⚠️ Lỗi lưu file lần {retry_count + 1}: {e}")
                if retry_count < max_retries - 1:
                    time.sleep(0.5)
        
        print(f"❌ Không thể lưu file sau {max_retries} lần thử")
        return None

def clean_json_string(text: str) -> str:
    if not text:
        return ""

    text = text.strip()

    # BƯỚC 1: Dùng Regex để bắt nội dung trong ```json ... ``` (nếu có)
    # re.DOTALL giúp dấu chấm (.) khớp với cả dòng mới (\n)
    # re.IGNORECASE để bắt cả ```JSON và ```json
    pattern = r"```(?:json)?(.*?)```"
    match = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
    
    if match:
        # Nếu tìm thấy markdown, lấy nội dung bên trong
        text = match.group(1).strip()
    
    # BƯỚC 2: "Săn" JSON bằng cách tìm dấu { đầu tiên và } cuối cùng
    # Bước này cực quan trọng cho môn Tự nhiên khi AI hay nói nhảm trước/sau JSON
    start_idx = text.find('{')
    end_idx = text.rfind('}')

    if start_idx != -1 and end_idx != -1 and end_idx > start_idx:
        # Cắt lấy đúng phần từ { đến }
        return text[start_idx : end_idx + 1]

    # Trường hợp tệ nhất: Trả về nguyên gốc để các hàm repair (AI sửa lỗi) xử lý tiếp
    return text

def repair_json_with_ai(broken_json_str: str, client) -> str:
    """Gửi JSON lỗi cho AI sửa"""
    print("⚠️ JSON lỗi. Đang yêu cầu AI sửa...")
    prompt_fix = f"""
Đoạn JSON sau bị lỗi cú pháp:

{broken_json_str}

NHIỆM VỤ:
1. Sửa lỗi cú pháp JSON (escape quotes, thêm phẩy, đóng ngoặc)
2. KHÔNG thay đổi nội dung Tiếng Việt
3. KHÔNG thay đổi công thức LaTeX (giữ nguyên \\frac, \\sqrt...)
4. CHỈ TRẢ VỀ JSON ĐÃ SỬA (không markdown, không giải thích)
    """
    repaired_text = client.send_data_to_check(prompt_fix)
    return clean_json_string(repaired_text)
def sanitize_latex_json(text: str) -> str:
    """
    Sanitize JSON chứa LaTeX một cách AN TOÀN
    
    Chiến lược:
    1. Chỉ xử lý BÊN TRONG chuỗi JSON (giữa dấu ngoặc kép)
    2. Giữ nguyên phần cấu trúc JSON (keys, colons, brackets)
    3. Escape backslash KHÔNG phải JSON escape hợp lệ
    """
    
    # Danh sách escape sequences hợp lệ trong JSON spec
    VALID_JSON_ESCAPES = {
        '\\\\', '\\"', '\\/', '\\b', '\\f', '\\n', '\\r', '\\t'
    }
    
    def fix_string_content(match):
        """
        Xử lý nội dung BÊN TRONG chuỗi JSON (giữa dấu ngoặc kép)
        match.group(0) = toàn bộ "..." (có dấu ")
        match.group(1) = nội dung giữa dấu " (không có dấu ")
        """
        full_match = match.group(0)
        content = match.group(1)
        
        # Nếu chuỗi rỗng, giữ nguyên
        if not content:
            return full_match
        
        result = []
        i = 0
        
        while i < len(content):
            char = content[i]
            
            if char == '\\':
                # Kiểm tra có phải escape hợp lệ không
                if i + 1 < len(content):
                    next_char = content[i + 1]
                    two_chars = char + next_char
                    
                    # Trường hợp 1: JSON escape hợp lệ (\\, \", \n, \t...)
                    if two_chars in VALID_JSON_ESCAPES:
                        result.append(two_chars)
                        i += 2
                        continue
                    
                    # Trường hợp 2: Unicode escape (\uXXXX)
                    if next_char == 'u' and i + 5 < len(content):
                        hex_part = content[i+2:i+6]
                        if len(hex_part) == 4 and all(c in '0123456789ABCDEFabcdef' for c in hex_part):
                            result.append(content[i:i+6])  # \uXXXX
                            i += 6
                            continue
                    
                    # Trường hợp 3: LaTeX command (VD: \frac, \sqrt, \sin)
                    # → Escape thành \\
                    result.append('\\\\')
                    i += 1
                else:
                    # Backslash ở cuối chuỗi → Escape
                    result.append('\\\\')
                    i += 1
            else:
                result.append(char)
                i += 1
        
        # Trả về chuỗi đã fix (VẪN CÓ dấu ngoặc kép)
        return '"' + ''.join(result) + '"'
    
    # Regex tìm tất cả chuỗi JSON: "..."
    # (?:[^"\\]|\\.)* nghĩa là: (không phải " hoặc \) HOẶC (\ theo sau bất kỳ ký tự nào)
    string_pattern = r'"((?:[^"\\]|\\.)*)"'
    
    sanitized = re.sub(string_pattern, fix_string_content, text)
    
    return sanitized

def parse_json_safely(json_str: str, client) -> Optional[Dict]:
    """Parse JSON an toàn với Sanitization và Retry AI"""
    # 1. Clean markdown
    cleaned_str = clean_json_string(json_str)
    
    # 2. Bước quan trọng: Sanitize LaTeX backslashes bằng thuật toán (nhanh và chính xác hơn AI)
    sanitized_str = sanitize_latex_json(cleaned_str)
    
    # Thử parse lần 1 (với chuỗi đã sanitize)
    try:
        return json.loads(sanitized_str, strict=False)
    except json.JSONDecodeError as e:
        print(f"❌ Lỗi JSON lần 1 (Logic): {e}")
        # Debug: In ra đoạn lỗi để kiểm tra nếu cần
        start = max(0, e.pos - 20)
        end = min(len(sanitized_str), e.pos + 20)
        print(f"Context: ...{sanitized_str[start:end]}...")
    
    # Thử sửa bằng AI (Fallback cuối cùng)
    try:
        # Lưu ý: Gửi chuỗi gốc (cleaned_str) hoặc chuỗi đã sanitize tùy chiến lược. 
        # Thường gửi chuỗi gốc để AI tự định dạng lại từ đầu sẽ an toàn hơn về ngữ nghĩa.
        repaired_str = repair_json_with_ai(cleaned_str, client)
        
        # Sau khi AI sửa, vẫn nên sanitize lại một lần nữa để chắc chắn
        repaired_str = sanitize_latex_json(repaired_str)
        
        return json.loads(repaired_str, strict=False)
    except json.JSONDecodeError as e:
        print(f"❌ Lỗi JSON lần 2 (AI Give up): {e}")
        return None
def generate_or_get_image(hinh_anh_data: Dict) -> tuple:
    """
    Xử lý gọi hàm sinh ảnh.
    Returns: (image_bytes, placeholder_text) - image_bytes là 1 object duy nhất
    """
    mo_ta = hinh_anh_data.get("mo_ta", hinh_anh_data.get("description", ""))
    mo_ta = str(mo_ta).strip()
    loai = hinh_anh_data.get("loai", "tu_mo_ta")
    
    if loai == "tu_mo_ta" and mo_ta:
        try:
            from modules.common.text2Image import generate_image_from_text
            # Hàm này trả về 1 bytes object (hoặc None)
            image_bytes = generate_image_from_text(mo_ta)
            if image_bytes:
                return image_bytes, None
            else:
                # Nếu API trả về None (do lỗi mạng hoặc quota)
                return None, f"⚠️ [Lỗi sinh ảnh] Server không trả về ảnh cho mô tả: {mo_ta}"
        except Exception as e:
            print(f"❌ Lỗi sinh ảnh: {e}")
            return None, f"⚠️ [Lỗi Code] {str(e)}"
    
    placeholder = f"🖼️ [Cần chèn hình: {mo_ta}]"
    return None, placeholder

def insert_image_or_placeholder(doc: Document, hinh_anh_data: Dict):
    """Chèn ảnh hoặc placeholder vào document"""
    image_bytes, placeholder = generate_or_get_image(hinh_anh_data)
    
    if image_bytes:
        try:
            image_stream = BytesIO(image_bytes)
            doc.add_picture(image_stream, width=Inches(4))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            print(f"❌ Lỗi chèn ảnh: {e}")
            p = doc.add_paragraph()
            run = p.add_run(f"⚠️ [Lỗi chèn ảnh: {str(e)}]")
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.italic = True
    
    elif placeholder:
        p = doc.add_paragraph()
        run = p.add_run(placeholder)
        run.font.color.rgb = RGBColor(200, 0, 0)
        run.italic = True
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return doc

import json
from typing import Dict, List, Optional, Any

class PromptBuilder:
    """
    Builder tạo prompt động, tối ưu riêng cho KHOA HỌC XÃ HỘI để tránh lỗi LaTeX
    Hỗ trợ: Trắc nghiệm 4 đáp án, Đúng/Sai, Trả lời ngắn.
    """
    
    @staticmethod
    def build_json_structure_hint(question_type: str) -> str:
        """
        Tạo hint cấu trúc JSON cho các loại đề
        """
        if question_type == "trac_nghiem_4_dap_an":
            return """
{
  "loai_de": "trac_nghiem_4_dap_an",
  "tong_so_cau": 80,
  "cau_hoi": [
    {
      "stt": 1,
      "muc_do": "nhan_biet",
      "phan": "PHẦN I: LỊCH SỬ THẾ GIỚI",
      "noi_dung": "Nội dung câu hỏi (Không dùng dấu $)...",
      "trich_dan": "Đoạn văn bản gốc từ PDF (Điền nếu có tư liệu, nếu không để chuỗi rỗng \"\")",
      "nguon_trich_dan": "(SGK Lịch sử 12, bộ Cánh Diều, trang 10) (Điền nếu có tư liệu)",
      "hinh_anh": {
        "co_hinh": true,
        "loai": "tu_mo_ta",
        "mo_ta": "Mô tả chi tiết hình ảnh (Bản đồ/Lược đồ/Chân dung)..."
      },
      "dap_an": [
        {"ky_hieu": "A", "noi_dung": "Nội dung đáp án A"},
        {"ky_hieu": "B", "noi_dung": "Nội dung đáp án B"},
        {"ky_hieu": "C", "noi_dung": "Nội dung đáp án C"},
        {"ky_hieu": "D", "noi_dung": "Nội dung đáp án D"}
      ],
      "dap_an_dung": 2,
      "giai_thich": "Giải thích chi tiết..."
    }
  ]
}
"""
        elif question_type == "dung_sai":
            return """
{
  "loai_de": "dung_sai",
  "tong_so_cau": 40,
  "cau_hoi": [
    {
      "stt": 1,
      "muc_do": "thong_hieu",
      "phan": "PHẦN I",
      "doan_thong_tin": "Đoạn tư liệu đầu bài...",
      "trich_dan": "Trích dẫn nguyên văn từ PDF (Điền nếu có tư liệu, nếu không để chuỗi rỗng \"\")",
      "nguon_trich_dan": "(Nguồn...) (Điền nếu có tư liệu)",
      "hinh_anh": { "co_hinh": false },
      "cac_y": [
        {"ky_hieu": "a", "noi_dung": "Phát biểu a", "dung": false},
        {"ky_hieu": "b", "noi_dung": "Phát biểu b", "dung": true},
        {"ky_hieu": "c", "noi_dung": "Phát biểu c", "dung": false},
        {"ky_hieu": "d", "noi_dung": "Phát biểu d", "dung": true}
      ],
      "dap_an_dung_sai": "0101",
      "giai_thich": [
        {"y": "a", "noi_dung_y": "...", "ket_luan": "SAI", "giai_thich": "Giải thích..."},
        {"y": "b", "noi_dung_y": "...", "ket_luan": "ĐÚNG", "giai_thich": "Giải thích..."}
      ]
    }
  ]
}
"""
        elif question_type == "tra_loi_ngan":
            return """
{
  "loai_de": "tra_loi_ngan",
  "tong_so_cau": 30,
  "cau_hoi": [
    {
      "stt": 1,
      "muc_do": "nhan_biet",
      "phan": "PHẦN I",
      "noi_dung": "Nội dung câu hỏi ngắn gọn...",
      "trich_dan": "Đoạn văn bản gốc (Điền nếu có tư liệu, nếu không để chuỗi rỗng \"\")",
      "nguon_trich_dan": "(Nguồn...) (Điền nếu có tư liệu)",
      "hinh_anh": {
        "co_hinh": true,
        "loai": "tu_mo_ta",
        "mo_ta": "Mô tả chi tiết hình ảnh..."
      },
      "dap_an": "Đáp án ngắn gọn (Vd: 1945 hoặc Hà Nội)",
      "giai_thich": "Giải thích chi tiết về đáp án..."
    }
  ]
}
"""
        return "{}"
    
    @staticmethod
    def wrap_user_prompt(user_prompt: str, question_type: str) -> str:
        json_hint = PromptBuilder.build_json_structure_hint(question_type)
        
        # PROMPT ĐƯỢC TỐI ƯU RIÊNG CHO XÃ HỘI (Chặn LaTeX, Citations linh hoạt)
        return f"""{user_prompt}

----------------
### HƯỚNG DẪN KỸ THUẬT (SYSTEM INSTRUCTION) - BẮT BUỘC TUÂN THỦ:

1. **QUY TẮC ĐỊNH DẠNG VĂN BẢN (KHẮC PHỤC LỖI LATEX - Dành cho Sử/Địa):**
   - Đây là môn **KHOA HỌC XÃ HỘI (Lịch Sử / Địa Lý)**.
   - **TUYỆT ĐỐI KHÔNG** sử dụng dấu `$` hoặc cặp dấu `$$` trong nội dung JSON.
   - **Xử lý Toạ độ Địa lý / Độ C / Phần trăm:**
     + SAI (Cấm): `$20^0N$`, `$105^0E$`, `$30^0C$`, `$25\%$`.
     + ĐÚNG (Bắt buộc): Dùng ký tự Unicode hoặc chữ thường.
       -> "20°B", "105°Đ", "20 độ Vĩ Bắc", "30°C", "25%".
   - **Xử lý Ngày tháng / Thế kỷ:**
     + SAI: `$thế kỷ XX$`, `$năm 1945$`.
     + ĐÚNG: "thế kỷ XX", "năm 1945".

2. **QUY TẮC TRÍCH DẪN TƯ LIỆU (CỨNG NHƯNG LINH HOẠT):**
   - **Yêu cầu BẮT BUỘC:** Nếu câu hỏi **dựa trên hoặc tham chiếu** đến một đoạn văn/tư liệu cụ thể (VD: Câu hỏi đọc hiểu trong Lịch sử, hoặc phân tích bảng số liệu trong Địa lý), bạn **PHẢI** điền đầy đủ 2 trường sau:
     + `"trich_dan"`: Copy nguyên văn đoạn text/tư liệu từ PDF làm căn cứ.
     + `"nguon_trich_dan"`: Ghi rõ nguồn gốc (SGK..., trang...).
   - **Trường hợp MIỄN TRỪ:** Nếu câu hỏi mang tính khái quát (VD: Tác dụng của gió mùa, năm diễn ra sự kiện...) hoặc chỉ dùng hình ảnh (Bản đồ) mà không cần trích dẫn nguyên văn văn bản, hãy để **chuỗi rỗng** `""` cho cả hai trường này.

3. **YÊU CẦU VỀ HÌNH ẢNH:**
   - Nếu nội dung liên quan đến Bản đồ, Lược đồ, Biểu đồ... -> BẮT BUỘC đặt `"co_hinh": true` và điền mô tả chi tiết vào `"mo_ta"`.

4. **FORMAT JSON OUTPUT:**
   - Chỉ trả về duy nhất 1 chuỗi JSON.
   - Không được có text dẫn nhập hay kết thúc.
   - Đảm bảo JSON valid.

### MẪU JSON MONG MUỐN:
{json_hint}
"""

# ============================================================================
# PHẦN 5: DYNAMIC DOCX RENDERER (MỚI - AUTO-ADAPT)
# ============================================================================

class DynamicDocxRenderer:
    def __init__(self, doc: Document):
        self.doc = doc
    
    def render_title(self, data: Dict):
        """Render tiêu đề tự động"""
        loai_de = data.get("loai_de", "").upper()
        title = self.doc.add_heading(f'ĐỀ {loai_de}', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def auto_group_questions(self, data: Dict) -> Dict[str, List]:
        """
        Tự động nhóm câu hỏi và CHUẨN HÓA key muc_do từ Tiếng Việt sang code.
        Giúp người dùng thoải mái viết prompt "Vận dụng", "Nhận biết"... mà không bị lỗi file trắng.
        """
        grouped = {}
        for cau in data.get("cau_hoi", []):
            # 1. Lấy dữ liệu thô từ AI (ví dụ: "Vận dụng", "Nhận biết", "Thông hiểu")
            # Chuyển về chữ thường để dễ so sánh
            raw_muc_do = str(cau.get("muc_do", "unknown")).lower().strip()
            
            # 2. Logic "Phiên dịch" thông minh (Mapping)
            # Ưu tiên check "cao" trước để phân biệt "Vận dụng" và "Vận dụng cao"
            if "cao" in raw_muc_do:
                muc_do_chuan = "van_dung_cao"
            elif "dụng" in raw_muc_do or "dung" in raw_muc_do:
                muc_do_chuan = "van_dung"
            elif "thông" in raw_muc_do or "thong" in raw_muc_do:
                muc_do_chuan = "thong_hieu"
            elif "nhận" in raw_muc_do or "nhan" in raw_muc_do:
                muc_do_chuan = "nhan_biet"
            else:
                # Trường hợp AI ghi nội dung lạ, mặc định đưa vào Vận dụng 
                # để đảm bảo câu hỏi vẫn hiện ra trong file (tránh lỗi trang trắng)
                muc_do_chuan = "van_dung" 
            
            # 3. Gom nhóm theo key chuẩn
            if muc_do_chuan not in grouped:
                grouped[muc_do_chuan] = []
            grouped[muc_do_chuan].append(cau)
        
        # Sắp xếp theo STT trong mỗi nhóm
        for key in grouped:
            grouped[key].sort(key=lambda x: x.get("stt", 0))
        
        return grouped
    
    def get_section_title(self, muc_do: str) -> str:
        """
        Tạo tiêu đề section dựa trên mức độ
        CÓ THỂ mở rộng bằng config file
        """
        mapping = {
            "nhan_biet": "I. CÂU HỎI NHẬN BIẾT",
            "thong_hieu": "II. CÂU HỎI THÔNG HIỂU",
            "van_dung": "III. CÂU HỎI VẬN DỤNG",
            "van_dung_cao": "IV. CÂU HỎI VẬN DỤNG CAO"
        }
        return mapping.get(muc_do, muc_do.upper())
    
    def render_question_trac_nghiem(self, cau: Dict):
        """Render câu hỏi trắc nghiệm 4 đáp án"""
        # Câu hỏi
        p = self.doc.add_paragraph()
        p.add_run(f"Câu {cau['stt']}. ").bold = True
        process_text_with_latex(cau['noi_dung'], p)
        
        # Hình ảnh
        hinh_anh = cau.get("hinh_anh", {})
        if hinh_anh.get("co_hinh"):
            insert_image_or_placeholder(self.doc, hinh_anh)
        
        # Đáp án - THÊM XỬ LÝ LATEX
        for dap_an in cau.get("dap_an", []):
            p_da = self.doc.add_paragraph()
            run_ky_hieu = p_da.add_run(f"{dap_an['ky_hieu']}. ")
            process_text_with_latex(dap_an['noi_dung'], p_da) 
        
        # Lời giải
        p_lg = self.doc.add_paragraph()
        p_lg.add_run("Lời giải").bold = True
        
        if "dap_an_dung" in cau:
            p_dung = self.doc.add_paragraph()
            p_dung.add_run(f"{cau['dap_an_dung']}").bold = True
            self.doc.add_paragraph("####")
        
        # Giải thích - THÊM XỬ LÝ LATEX
        giai_thich = cau.get("giai_thich", "")
        for line in giai_thich.split("\n"):
            if line.strip():
                p_gt = self.doc.add_paragraph()
                process_text_with_latex(line.strip(), p_gt)  
        
        # Kết luận - THÊM XỬ LÝ LATEX
        if "dap_an_dung" in cau:
            dap_an_num = cau['dap_an_dung']
            noi_dung_dap_an = cau['dap_an'][dap_an_num-1]['noi_dung']
            p_ket_luan = self.doc.add_paragraph()
            run = p_ket_luan.add_run("Vậy đáp án đúng là: ")
            run.bold = True
            process_text_with_latex(noi_dung_dap_an, p_ket_luan, bold=True) 
    
    def render_question_dung_sai(self, cau: Dict):
        """Render câu hỏi đúng/sai"""
        # Số câu
        p = self.doc.add_paragraph()
        p.add_run(f"Câu {cau['stt']}.").bold = True
        
        # Đoạn thông tin - THÊM XỬ LÝ LATEX
        if cau.get("doan_thong_tin"):
            p_doan = self.doc.add_paragraph()
            process_text_with_latex(cau.get("doan_thong_tin", ""), p_doan)  
        
        # Hình ảnh
        hinh_anh = cau.get("hinh_anh", {})
        if hinh_anh.get("co_hinh"):
            insert_image_or_placeholder(self.doc, hinh_anh)
        
        # Các ý a, b, c, d - THÊM XỬ LÝ LATEX
        for y in cau.get("cac_y", []):
            p_y = self.doc.add_paragraph()
            p_y.add_run(f"{y['ky_hieu']}) ")
            process_text_with_latex(y['noi_dung'], p_y)  
        
        # Lời giải
        p_lg = self.doc.add_paragraph()
        p_lg.add_run("Lời giải").bold = True
        
        p_da = self.doc.add_paragraph()
        p_da.add_run(cau.get("dap_an_dung_sai", "")).bold = True
        self.doc.add_paragraph("####")
        
        # Giải thích từng ý - THÊM XỬ LÝ LATEX
        for gt in cau.get("giai_thich", []):
            p_gt = self.doc.add_paragraph()
            p_gt.add_run('+) "')
            process_text_with_latex(gt.get('noi_dung_y', ''), p_gt)  
            run_kl = p_gt.add_run(f'" - {gt.get("ket_luan", "SAI")}. ')
            run_kl.bold = True
            
            if gt.get('giai_thich'):
                # p_gt_detail = self.doc.add_paragraph()
                process_text_with_latex(gt.get('giai_thich', ''), p_gt)  
    
    def render_question_tra_loi_ngan(self, cau: Dict):
        """Render câu hỏi trả lời ngắn"""
        # Câu hỏi
        p = self.doc.add_paragraph()
        p.add_run(f"Câu {cau['stt']}. ").bold = True
        p_noi_dung = self.doc.add_paragraph()
        process_text_with_latex(cau['noi_dung'], p_noi_dung)  
        
        # Hình ảnh (nếu có)
        hinh_anh = cau.get("hinh_anh", {})
        if hinh_anh.get("co_hinh"):
            insert_image_or_placeholder(self.doc, hinh_anh)
        
        # Đáp án - THÊM XỬ LÝ LATEX
        p_da = self.doc.add_paragraph()
        run_label = p_da.add_run("Đáp án: ")
        run_label.bold = True
        
        raw_ans = str(cau.get('dap_an', '')).strip()
        if raw_ans.startswith("[[") and raw_ans.endswith("]]"):
            final_ans = raw_ans
        else:
            final_ans = f"[[{raw_ans}]]"
        
        # XỬ LÝ LATEX TRONG ĐÁP ÁN
        process_text_with_latex(final_ans, p_da, bold=True)  
        
        # Lời giải header
        p_lg = self.doc.add_paragraph()
        p_lg.add_run("Lời giải").bold = True
        self.doc.add_paragraph("####")
        
        # Giải thích chi tiết - ĐÃ CÓ XỬ LÝ LATEX
        giai_thich = cau.get("giai_thich", "")
        lines = giai_thich.replace('\\n', '\n').split('\n')
        
        for line in lines:
            text = line.strip()
            if not text or text == "####":
                continue
            
            is_bold = False
            if text.startswith("**") and text.endswith("**"):
                text = text[2:-2]
                is_bold = True
            
            check_text = text.replace('*', '').strip().lower()
            if check_text.startswith("vậy"):
                is_bold = True
                text = text.replace('**', '')

            p_gt = self.doc.add_paragraph()
            process_text_with_latex(text, p_gt, bold=is_bold)  
    
    def render_all(self, data: Dict):
        """
        Main render function - Có hỗ trợ chia PHẦN (PART) bên trong Mức độ
        """
        self.render_title(data)
        
        # 1. Auto-group theo mức độ (Nhận biết, Thông hiểu...)
        grouped = self.auto_group_questions(data)
        
        # 2. Detect loại đề
        loai_de = data.get("loai_de", "")
        
        # 3. Render từng nhóm MỨC ĐỘ
        # Thứ tự ưu tiên render
        order_muc_do = ["nhan_biet", "thong_hieu", "van_dung", "van_dung_cao"]
        
        for muc_do in order_muc_do:
            if muc_do not in grouped:
                continue
            
            # Lấy danh sách câu hỏi trong mức độ này
            questions = grouped[muc_do]
            if not questions:
                continue
            section_title = self.get_section_title(muc_do)
            self.doc.add_heading(section_title, level=2)
            current_phan = None

            for cau in questions:
                # Lấy tên phần của câu hiện tại
                phan_cua_cau = str(cau.get("phan", "")).strip()
                
                # Nếu câu này thuộc một phần mới -> In Header Phần
                if phan_cua_cau and phan_cua_cau != current_phan:
                    # In ra header cấp 3 (VD: Phần 1: Đội ngũ...)
                    # Dùng màu hoặc in đậm để phân biệt
                    p_phan = self.doc.add_heading(phan_cua_cau.upper(), level=3)
                    p_phan.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    current_phan = phan_cua_cau
                
                # Render nội dung câu hỏi như bình thường
                if loai_de == "dung_sai":
                    self.render_question_dung_sai(cau)
                elif loai_de == "tra_loi_ngan":
                    self.render_question_tra_loi_ngan(cau)
                else:
                    self.render_question_trac_nghiem(cau)

def response2docx_flexible(
    file_path: str,
    prompt: str,
    file_name: str,
    project_id: str,
    creds: str,
    model_name: str,
    question_type: str = "trac_nghiem_4_dap_an",
    batch_name: Optional[str] = None
) -> Optional[str]:
    try:
        from modules.common.callAPI import VertexClient
        
        client = VertexClient(project_id, creds, model_name)
        
        if not batch_name:
            batch_name = file_name.replace("_TN", "").replace("_DS", "").replace("_TLN", "")
        
        # 1. Wrap prompt với JSON structure hint
        final_prompt = PromptBuilder.wrap_user_prompt(prompt, question_type)
        
        # 2. Gửi request AI
        print("📤 Đang gửi request tới AI...")
        ai_response = client.send_data_to_AI(final_prompt, file_path)
        
        # 3. Parse JSON
        print("🔄 Đang parse JSON...")
        data = parse_json_safely(ai_response, client)
        if not data:
            print("❌ Không thể parse JSON từ AI")
            return None
        
        print(f"✅ Parse thành công: {data.get('tong_so_cau', 0)} câu hỏi")
        
        # 4. Render DOCX động
        print("📝 Đang tạo DOCX...")
        doc = Document()
        renderer = DynamicDocxRenderer(doc)
        
        try:
            renderer.render_all(data)
            print("✅ Render DOCX thành công")
        except Exception as e:
            print(f"❌ Lỗi khi render DOCX: {e}")
            traceback.print_exc()
            return None
        
        # 5. Lưu file
        print("💾 Đang lưu file...")
        output_path = save_document_securely(doc, batch_name, file_name)
        
        if output_path:
            print(f"✅ Hoàn thành: {output_path}")
        else:
            print("❌ Không thể lưu file")
            
        return output_path
    
    except Exception as e:
        print(f"❌ LỖI NGHIÊM TRỌNG: {e}")
        traceback.print_exc()
        return None

def response2docx_json(file_path, prompt, file_name, project_id, creds, model_name, batch_name=None):
    """Wrapper cho trắc nghiệm 4 đáp án (legacy)"""
    return response2docx_flexible(
        file_path, prompt, file_name, project_id, creds, model_name,
        question_type="trac_nghiem_4_dap_an",
        batch_name=batch_name
    )

def response2docx_dung_sai_json(file_path, prompt, file_name, project_id, creds, model_name, batch_name=None):
    """Wrapper cho đúng/sai (legacy)"""
    return response2docx_flexible(
        file_path, prompt, file_name, project_id, creds, model_name,
        question_type="dung_sai",
        batch_name=batch_name
    )
    
def response2docx_tra_loi_ngan_json(file_path, prompt, file_name, project_id, creds, model_name, batch_name=None):
    """Wrapper cho trả lời ngắn (legacy compatibility)"""
    return response2docx_flexible(
        file_path, prompt, file_name, project_id, creds, model_name,
        question_type="tra_loi_ngan",
        batch_name=batch_name
    )

class ConfigManager:
    """
    Quản lý cấu hình qua file JSON
    Cho phép thay đổi TOÀN BỘ behavior mà không sửa code
    """
    
    DEFAULT_CONFIG = {
        "section_mapping": {
            "nhan_biet": "I. CÂU HỎI NHẬN BIẾT",
            "thong_hieu": "II. CÂU HỎI THÔNG HIỂU",
            "van_dung": "III. CÂU HỎI VẬN DỤNG",
            "van_dung_cao": "IV. CÂU HỎI VẬN DỤNG CAO" 
        },
        "section_order": ["nhan_biet", "thong_hieu", "van_dung", "van_dung_cao"],
        "auto_fix": True,
        "image_width_inches": 4,
        "retry_json_parse": 2
    }
    
    @classmethod
    def load_config(cls, config_path: str = "config.json") -> Dict:
        """Load config từ file hoặc dùng default"""
        if os.path.exists(config_path):
            with open(config_path, "r", encoding="utf-8") as f:
                return json.load(f)
        return cls.DEFAULT_CONFIG
    
    @classmethod
    def save_config(cls, config: Dict, config_path: str = "config.json"):
        """Lưu config để tái sử dụng"""
        with open(config_path, "w", encoding="utf-8") as f:
            json.dump(config, f, ensure_ascii=False, indent=2)

