import json
import os
import sys
import threading
import time
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from typing import Dict, List, Optional, Any
import zipfile
import subprocess
import re
from tempfile import NamedTemporaryFile
from docx.oxml import parse_xml
import traceback
from modules.common.schema import (
    schema_trac_nghiem, 
    schema_dung_sai, 
    schema_tra_loi_ngan, 
    schema_tu_luan
)

_FILE_LOCK = threading.RLock()
_OUTPUT_DIR_LOCK = threading.RLock()

def get_app_path():
    """Lấy đường dẫn chứa file .exe hoặc script"""
    if getattr(sys, 'frozen', False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.abspath(__file__))
def sanitize_xml_string(text):
    """
    Loại bỏ các ký tự điều khiển không hợp lệ trong XML (ASCII 0-31, trừ 9, 10, 13).
    """
    if not text:
        return ""
    # Regex loại bỏ các ký tự từ \x00-\x08, \x0B-\x0C, \x0E-\x1F
    return re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F]', '', str(text))
def find_pandoc_executable():
    """
    Tìm pandoc.exe theo thứ tự ưu tiên:
    1. Thư mục 'pandoc' cạnh tool (cho bản build)
    2. PATH hệ thống (cho môi trường dev)
    """
    app_path = get_app_path()
    
    # 1. Tìm trong thư mục cục bộ 'pandoc' (ưu tiên cao nhất)
    local_pandoc = os.path.join(app_path, 'pandoc', 'pandoc.exe')
    if os.path.isfile(local_pandoc):
        # print(f"✅ Sử dụng Pandoc cục bộ: {local_pandoc}")
        return local_pandoc
    
    # 2. Fallback: Tìm trong PATH hệ thống (cho dev)
    import shutil
    system_pandoc = shutil.which('pandoc')
    if system_pandoc:
        # print(f"⚠️ Sử dụng Pandoc hệ thống: {system_pandoc}")
        return system_pandoc
    
    # 3. Không tìm thấy
    print("❌ KHÔNG TÌM THẤY PANDOC!")
    return None

def latex_to_omml_via_pandoc(latex_math_dollar):
    """Chuyển đổi LaTeX sang OMML qua Pandoc"""
    pandoc_exe = find_pandoc_executable()
    
    if not pandoc_exe:
        print("❌ Pandoc không khả dụng, bỏ qua equation")
        return None
    
    try:
        # Chuẩn hóa input (loại bỏ ký tự lạ)
        latex_clean = latex_math_dollar.strip()
        
        # Tạo file tạm với encoding UTF-8 BOM để tránh lỗi
        with NamedTemporaryFile(mode='w', suffix=".docx", delete=False, encoding='utf-8') as temp_docx:
            temp_path = temp_docx.name
        
        # Chạy Pandoc với error handling tốt hơn
        result = subprocess.run(
            [pandoc_exe, '--from=latex', '--to=docx', '-o', temp_path],
            input=latex_clean,
            text=True,
            encoding='utf-8',
            capture_output=True,
            timeout=10,  # Timeout 10s để tránh treo
            creationflags=subprocess.CREATE_NO_WINDOW if hasattr(subprocess, 'CREATE_NO_WINDOW') else 0
        )
 
        if result.returncode != 0:
            error_msg = result.stderr.strip() if result.stderr else "Unknown error"
            # print(f"⚠️ Pandoc error (code {result.returncode}): {error_msg}")
            
            # Kiểm tra lỗi phổ biến
            if "not found" in error_msg.lower() or "cannot find" in error_msg.lower():
                print("   → Thiếu DLL dependencies. Kiểm tra lại folder pandoc/")
            elif "syntax" in error_msg.lower():
                print(f"   → LaTeX syntax error: {latex_clean[:50]}...")
            
            return None
        
        # Kiểm tra file output có tồn tại không
        if not os.path.exists(temp_path) or os.path.getsize(temp_path) == 0:
            # print(f"⚠️ Pandoc không tạo file output hợp lệ")
            return None
           
        # Đọc XML từ DOCX
        with zipfile.ZipFile(temp_path, 'r') as z:
            xml_content = z.read('word/document.xml').decode('utf-8')
        
        # Dọn dẹp file tạm
        try:
            os.remove(temp_path)
        except:
            pass
       
        # Tìm equation XML
        match = re.search(r'(<m:oMath[^>]*>.*?</m:oMath>)', xml_content, re.DOTALL)
        
        if not match:
            # print(f"⚠️ Không tìm thấy equation trong output: {latex_clean[:30]}...")
            return None
            
        return match.group(1)
   
    except subprocess.TimeoutExpired:
        print(f"⚠️ Pandoc timeout (>10s)")
        return None
    except Exception as e:
        print(f"❌ Lỗi latex_to_omml: {type(e).__name__}: {e}")
        import traceback
        traceback.print_exc()
        return None



def process_text_with_latex(text, paragraph, bold=False):
    """
    Xử lý text có công thức LaTeX
    VERSION ỔN ĐỊNH - Copy từ test_res.py (KHÔNG có repair_broken_latex)
    """
    if not text:
        return
    text = sanitize_xml_string(text).strip()
    is_entirely_bold = bold
    if text.startswith("**") and text.endswith("**"):
        is_entirely_bold = True
        text = text[2:-2]
    # Làm sạch HTML tags
    text = text.replace("<br>", "\n").replace("<br/>", "\n") \
               .replace("<Br>", "\n").replace("<Br/>", "\n")
    text = re.sub(r'</?(div|p|u|span|font|i|b)\b[^>]*>', '', text)
    text = text.replace("&nbsp;", "").replace("&lt;", "").replace("&gt;", "")
    
    # Tách text và LaTeX
    pattern = r'(\$[^$]+\$|\\\[.*?\\\])'
    parts = re.split(pattern, text)
    
    for part in parts:
        if not part:
            continue
        
        # Phần LaTeX
        if part.startswith('$') or part.startswith('\\['):
            try:
                latex_expr = clean_latex_math(part)
                insert_equation_into_paragraph(latex_expr, paragraph)
            except Exception as e:
                # Fallback: thêm text thuần
                run = paragraph.add_run(part)
                run.bold = is_entirely_bold
        # Phần text thường
        else:
            # cleaned_part = re.sub(r'^\s*/', '', part)
            # run = paragraph.add_run(cleaned_part)
            # if bold:
            #     run.bold = True
            sub_parts = re.split(r'(\*\*.*?\*\*)', part)
            for sp in sub_parts:
                sp_clean = sanitize_xml_string(sp)
                if not sp_clean: 
                    continue
                if sp.startswith("**") and sp.endswith("**"):
                    run = paragraph.add_run(sp[2:-2])
                    run.bold = True
                else:
                    run = paragraph.add_run(sp)
                    run.bold = is_entirely_bold


def insert_equation_into_paragraph(latex_math_dollar, paragraph):
    """Chèn công thức toán học vào paragraph"""
    omml_str = latex_to_omml_via_pandoc(latex_math_dollar)
    
    if not omml_str:
        # Fallback: Thêm text thuần nếu không convert được
        paragraph.add_run(f" [{latex_math_dollar}] ")
        return
    
    # Thêm namespace nếu thiếu
    if 'xmlns:m=' not in omml_str:
        omml_str = re.sub(
            r'<m:oMath',
            r'<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"',
            omml_str,
            count=1
        )
    
    try:
        omml_element = parse_xml(omml_str)
        run = paragraph.add_run()
        run._r.append(omml_element)
    except Exception as e:
        print(f"Lỗi chèn equation: {e}")
        paragraph.add_run(f" [{latex_math_dollar}] ")


def clean_latex_math(latex_raw):
    latex_raw = latex_raw.lstrip('$').rstrip('$')
    latex_raw = latex_raw.strip()
    latex_raw = re.sub(r'\\/', '', latex_raw)
    latex_raw = re.sub(r'\\operatorname\s*{\s*([^}]*)\s*}',
                       lambda m: m.group(1).replace(' ', ''), latex_raw)
    latex_raw = re.sub(r'\\root\s*(\d+)\s*{([^}]*)}', r'\\sqrt[\1]{\2}', latex_raw)
    latex_raw = re.sub(r'\\root\s*{(\d+)}\s*\\of\s*{([^}]*)}', r'\\sqrt[\1]{\2}', latex_raw)
    latex_raw = re.sub(r'\\root\s*(\d+)\s*\\sqrt\s*{([^}]*)}', r'\\sqrt[\1]{\2}', latex_raw)
    latex_raw = re.sub(r'([a-zA-Z])\s*\\frac\s*{([^}]+)}\s*{([^}]+)}',
                       r'\1^{\\frac{\2}{\3}}', latex_raw)
    latex_raw = re.sub(r'\\sp\s*{([^}]*)}', r'^{\1}', latex_raw)
    latex_raw = re.sub(r'{\\bf\s*([^}]*)}', r'\1', latex_raw)
    latex_raw = re.sub(r'\\\s*log', r'\\log', latex_raw)
    latex_raw = re.sub(r'\\bigskip', '', latex_raw)
    latex_raw = re.sub(r'\\nonumber', '', latex_raw)
    latex_raw = latex_raw.replace(r'\?', '?')
    latex_raw = re.sub(r'\\cdot\s*(?=\w)', r'\\cdot ', latex_raw)
    latex_raw = latex_raw.replace(r'\dotstan', r'\cdot \tan')
    latex_raw = re.sub(r'(?<!\\)(\bln\b|\blog\b|\bsin\b|\bcos\b|\btan\b|\blog_{?\d*}?)',
                       r'\\\1', latex_raw)
    latex_raw = re.sub(r'(\\Leftrightarrow|\\Rightarrow|\\rightarrow)(?=\w)', r'\1 ', latex_raw)
    latex_raw = latex_raw.replace(r'\\n', r'\n')
    
    latex_raw = latex_raw.strip()
    # ✅ KHÁC BIỆT: Version cũ KHÔNG replace \n và \r
    # latex_raw = latex_raw.replace('\n', ' ').replace('\r', '')  # ← XÓA DÒNG NÀY
    
    if not (latex_raw.startswith('$') and latex_raw.endswith('$')):
        latex_raw = f"${latex_raw}$"
    
    return latex_raw

def ensure_output_folder_for_batch(batch_name):
    """Tạo folder riêng cho batch"""
    base_path = get_app_path()
    output_base = os.path.join(base_path, "output")
    batch_folder = os.path.join(output_base, batch_name)
    
    with _OUTPUT_DIR_LOCK:
        os.makedirs(output_base, exist_ok=True)
        os.makedirs(batch_folder, exist_ok=True)
    
    return batch_folder

def save_document_securely(doc, batch_name, file_name):
    """Lưu file DOCX với thread-safety"""
    batch_folder = ensure_output_folder_for_batch(batch_name)
    if not batch_folder:
        return None

    output_path = os.path.join(batch_folder, f"{file_name}.docx")
    
    with _FILE_LOCK:
        max_retries = 3
        for retry_count in range(max_retries):
            try:
                doc.save(output_path)
                if os.path.exists(output_path):
                    file_size = os.path.getsize(output_path)
                    print(f"✅ Đã lưu file: {output_path}")
                    return output_path
            except Exception as e:
                print(f"⚠️ Lỗi lưu file lần {retry_count + 1}: {e}")
                if retry_count < max_retries - 1:
                    time.sleep(1)
        
        print(f"❌ Không thể lưu file sau {max_retries} lần thử")
        return None
def save_json_securely(data, batch_name, file_name):
    """Lưu file JSON với thread-safety"""
    batch_folder = ensure_output_folder_for_batch(batch_name)
    if not batch_folder: return None

    output_path = os.path.join(batch_folder, f"{file_name}.json")
    with _FILE_LOCK:
        try:
            with open(output_path, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            print(f"✅ Đã lưu file JSON: {output_path}")
            return output_path
        except Exception as e:
            print(f"❌ Lỗi lưu file JSON: {e}")
            return None
def generate_or_get_image(hinh_anh_data: Dict) -> tuple:
    mo_ta = hinh_anh_data.get("mo_ta", hinh_anh_data.get("description", ""))
    mo_ta = str(mo_ta).strip()
    loai = hinh_anh_data.get("loai", "tu_mo_ta")
    
    if loai == "tu_mo_ta" and mo_ta:
        try:
            from modules.common.text2Image import generate_image_from_text
            # Hàm này trả về 1 bytes object (hoặc None)
            image_bytes = generate_image_from_text(mo_ta)
            if image_bytes:
                return image_bytes, None
            else:
                # Nếu API trả về None (do lỗi mạng hoặc quota)
                return None, f"⚠️ [Lỗi sinh ảnh] Server không trả về ảnh cho mô tả: {mo_ta}"
        except Exception as e:
            print(f"❌ Lỗi sinh ảnh: {e}")
            return None, f"⚠️ [Lỗi Code] {str(e)}"
    
    placeholder = f"🖼️ [Cần chèn hình: {mo_ta}]"
    return None, placeholder

def insert_image_or_placeholder(doc: Document, hinh_anh_data: Dict):
    """Chèn ảnh hoặc placeholder vào document"""
    image_bytes, placeholder = generate_or_get_image(hinh_anh_data)
    
    if image_bytes:
        try:
            image_stream = BytesIO(image_bytes)
            doc.add_picture(image_stream, width=Inches(4))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            print(f"❌ Lỗi chèn ảnh: {e}")
            p = doc.add_paragraph()
            run = p.add_run(f"⚠️ [Lỗi chèn ảnh: {str(e)}]")
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.italic = True
    
    elif placeholder:
        p = doc.add_paragraph()
        run = p.add_run(placeholder)
        run.font.color.rgb = RGBColor(200, 0, 0)
        run.italic = True
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return doc

class PromptBuilder:
    """
    PromptBuilder mới: Gọn nhẹ, chỉ tập trung vào nội dung.
    Không còn nhồi nhét cấu trúc JSON vào prompt.
    """
    @staticmethod
    def wrap_user_prompt(user_prompt: str) -> str:
        # Chỉ giữ lại các quy tắc cốt lõi về nội dung
        return f"""{user_prompt}

----------------
### YÊU CẦU KỸ THUẬT QUAN TRỌNG:

1. **TUÂN THỦ SCHEMA:**
   - Output trả về PHẢI khớp chính xác với cấu trúc JSON Schema đã được định nghĩa trong hệ thống.
   - Tuyệt đối không thêm lời dẫn, không thêm markdown (```json). Chỉ trả về Raw JSON.

2. **QUY TẮC LATEX (BẮT BUỘC):**
   - Mọi công thức Toán/Lý/Hóa phải đặt trong dấu `$`.
   - Ví dụ: $x^2 + 2x$, $H_2SO_4$.
   - **Lưu ý escape:** Trong chuỗi JSON, ký tự backslash `\\` phải được nhân đôi thành `\\\\`. 
     Ví dụ: muốn viết $\\frac{{1}}{{2}}$ thì trong JSON phải là "$\\\\frac{{1}}{{2}}$".

3. **HÌNH ẢNH:**
   - Luôn điền trường "mo_ta" chi tiết nếu câu hỏi cần hình minh họa (đồ thị, thí nghiệm, bản đồ...).
"""

def get_schema_by_type(question_type: str):
    """Mapping loại câu hỏi sang Schema object"""
    mapping = {
        "trac_nghiem_4_dap_an": schema_trac_nghiem,
        "dung_sai": schema_dung_sai,
        "tra_loi_ngan": schema_tra_loi_ngan,
        "tu_luan": schema_tu_luan
    }
    return mapping.get(question_type, schema_trac_nghiem)
class DynamicDocxRenderer:
    """
    Renderer tự động thích ứng với cấu trúc JSON
    """
    
    def __init__(self, doc: Document):
        self.doc = doc
    
    def render_title(self, data: Dict):
        """Render tiêu đề tự động"""
        loai_de = data.get("loai_de", "").upper()
        title = self.doc.add_heading(f'ĐỀ {loai_de}', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ma_bai = data.get("ma_bai", "")
        if ma_bai:
            p_ma_bai = self.doc.add_paragraph()
            p_ma_bai.add_run(f"[{ma_bai},,]")
    def render_ma_dang_header(self, ma_dang: str):
        """Hiển thị [ma_dang] trên một dòng riêng biệt"""
        if ma_dang:
            p_ma = self.doc.add_paragraph()
            p_ma.add_run(f"[{ma_dang},,]")
    def auto_group_questions(self, data: Dict) -> Dict[str, List]:
        """
        Tự động nhóm câu hỏi và CHUẨN HÓA key muc_do từ Tiếng Việt sang code.
        Giúp người dùng thoải mái viết prompt "Vận dụng", "Nhận biết"... mà không bị lỗi file trắng.
        """
        grouped = {}
        for cau in data.get("cau_hoi", []):
            # 1. Lấy dữ liệu thô từ AI (ví dụ: "Vận dụng", "Nhận biết", "Thông hiểu")
            # Chuyển về chữ thường để dễ so sánh
            raw_muc_do = str(cau.get("muc_do", "unknown")).lower().strip()
            
            # 2. Logic "Phiên dịch" thông minh (Mapping)
            # Ưu tiên check "cao" trước để phân biệt "Vận dụng" và "Vận dụng cao"
            if "cao" in raw_muc_do:
                muc_do_chuan = "van_dung_cao"
            elif "dụng" in raw_muc_do or "dung" in raw_muc_do:
                muc_do_chuan = "van_dung"
            elif "thông" in raw_muc_do or "thong" in raw_muc_do:
                muc_do_chuan = "thong_hieu"
            elif "nhận" in raw_muc_do or "nhan" in raw_muc_do:
                muc_do_chuan = "nhan_biet"
            else:
                # Trường hợp AI ghi nội dung lạ, mặc định đưa vào Vận dụng 
                # để đảm bảo câu hỏi vẫn hiện ra trong file (tránh lỗi trang trắng)
                muc_do_chuan = "van_dung" 
            
            # 3. Gom nhóm theo key chuẩn
            if muc_do_chuan not in grouped:
                grouped[muc_do_chuan] = []
            grouped[muc_do_chuan].append(cau)
        
        # Sắp xếp theo STT trong mỗi nhóm
        for key in grouped:
            grouped[key].sort(key=lambda x: x.get("stt", 0))
        
        return grouped
    
    def get_section_title(self, muc_do: str) -> str:
        """
        Tạo tiêu đề section dựa trên mức độ
        CÓ THỂ mở rộng bằng config file
        """
        mapping = {
            "nhan_biet": "I. CÂU HỎI NHẬN BIẾT",
            "thong_hieu": "II. CÂU HỎI THÔNG HIỂU",
            "van_dung": "III. CÂU HỎI VẬN DỤNG",
            "van_dung_cao": "IV. CÂU HỎI VẬN DỤNG CAO"
        }
        return mapping.get(muc_do, muc_do.upper())
    
    def render_question_trac_nghiem(self, cau: Dict):
        self.render_ma_dang_header(cau.get("ma_dang"))
        p = self.doc.add_paragraph()
        p.add_run(f"Câu {cau['stt']}. ").bold = True
        process_text_with_latex(cau.get('noi_dung', ''), p)
        
        hinh_anh = cau.get("hinh_anh", {})
        if hinh_anh.get("co_hinh"):
            insert_image_or_placeholder(self.doc, hinh_anh)
        
        for dap_an in cau.get("cac_lua_chon", []):
            p_da = self.doc.add_paragraph()
            p_da.add_run(f"{dap_an['ky_hieu']}. ").bold = True
            process_text_with_latex(dap_an.get('noi_dung', ''), p_da)

        if cau.get('noi_dung_en'):
            self.doc.add_paragraph("(translate_en)").italic = True
            p_en = self.doc.add_paragraph()
            process_text_with_latex(cau.get('noi_dung_en', ''), p_en)
            
            for dap_an in cau.get("cac_lua_chon", []):
                p_da_en = self.doc.add_paragraph()
                p_da_en.add_run(f"{dap_an['ky_hieu']}. ").bold = True
                content_en = dap_an.get('noi_dung_en') or dap_an.get('noi_dung', '')
                process_text_with_latex(content_en, p_da_en)

        # --- PHẦN 2: LỜI GIẢI CHI TIẾT ---
        p_lg = self.doc.add_paragraph()
        p_lg.add_run("Lời giải").bold = True # <-- Header vẫn in đậm
        
        if "dap_an_dung" in cau:
            p_dung = self.doc.add_paragraph()
            p_dung.add_run(f"{cau['dap_an_dung']}").bold = True
            self.doc.add_paragraph("####") 

        # Hàm render từng dòng giải thích
        def render_explanation_lines(text_block, lang='vi'):
            if not text_block: return
            
            lines = text_block.split("\n")
            for line in lines:
                text = line.strip()
                if not text: continue
                
                p_gt = self.doc.add_paragraph()
                
                # --- LOGIC MỚI: CHỈ IN ĐẬM KẾT LUẬN ---
                is_bold = False
                text_lower = text.lower()
                
                if lang == 'vi':
                    # Chỉ in đậm dòng bắt đầu bằng "Vậy..."
                    if text_lower.startswith("vậy đáp án") or text_lower.startswith("vậy, đáp án"):
                        is_bold = True
                else: # English
                    # Chỉ in đậm dòng bắt đầu bằng "Therefore..."
                    if text_lower.startswith("therefore"):
                        is_bold = True
                
                process_text_with_latex(text, p_gt, bold=is_bold)

        # 2.1 Giải thích Tiếng Việt
        render_explanation_lines(cau.get("giai_thich", ""), lang='vi')

        # 2.2 Giải thích Tiếng Anh
        if cau.get("giai_thich_en"):
            self.doc.add_paragraph("(translate_en)").italic = True
            render_explanation_lines(cau.get("giai_thich_en", ""), lang='en')

        # --- PHẦN 3: GỢI Ý (Giữ nguyên) ---
        goi_y_vi = cau.get("goi_y", "")
        goi_y_en = cau.get("goi_y_en", "")
        
        if goi_y_vi or goi_y_en:
            self.doc.add_paragraph("####")
            if goi_y_vi:
                p_title = self.doc.add_paragraph()
                p_title.add_run("Gợi ý:").bold = True
                for line in goi_y_vi.split("\n"):
                    if not line.startswith("Gợi ý"):
                        process_text_with_latex(line.strip(), self.doc.add_paragraph())
            
            if goi_y_en:
                self.doc.add_paragraph("(translate_en)").italic = True
                p_title_en = self.doc.add_paragraph()
                p_title_en.add_run("Hint:").bold = True
                for line in goi_y_en.split("\n"):
                    if not line.startswith("Hint"):
                        process_text_with_latex(line.strip(), self.doc.add_paragraph())

    def render_question_dung_sai(self, cau: Dict):
        """
        Render câu hỏi Đúng/Sai với GIẢI THÍCH DẠNG MẢNG ĐỐI TƯỢNG
        """
        # 1. Render Header câu hỏi
        self.render_ma_dang_header(cau.get("ma_dang"))
        p = self.doc.add_paragraph()
        p.add_run(f"Câu {cau['stt']}. ").bold = True
       
        # 2. Render Đoạn thông tin ngữ cảnh
        if cau.get("doan_thong_tin"):
            process_text_with_latex(cau.get("doan_thong_tin", ""), p)
       
        # 3. Render Hình ảnh (nếu có)
        hinh_anh = cau.get("hinh_anh", {})
        if hinh_anh.get("co_hinh"):
            insert_image_or_placeholder(self.doc, hinh_anh)
       
        # 4. Render các ý a, b, c, d
        for y in cau.get("cac_y", []):
            p_y = self.doc.add_paragraph()
            p_y.add_run(f"{y['ky_hieu']}) ")
            process_text_with_latex(y.get('noi_dung', ''), p_y)
 
        # 5. Render Tiếng Anh (nếu có)
        has_en = cau.get("doan_thong_tin_en") or any(y.get('noi_dung_en') for y in cau.get("cac_y", []))
        if has_en:
            self.doc.add_paragraph("(translate_en)").italic = True
            if cau.get("doan_thong_tin_en"):
                p_en = self.doc.add_paragraph()
                process_text_with_latex(cau.get("doan_thong_tin_en", ""), p_en)
           
            for y in cau.get("cac_y", []):
                p_y_en = self.doc.add_paragraph()
                p_y_en.add_run(f"{y['ky_hieu']}) ")
                content_en = y.get('noi_dung_en') or y.get('noi_dung', '')
                process_text_with_latex(content_en, p_y_en)
 
        # --- PHẦN LỜI GIẢI (CẬP NHẬT: XỬ LÝ MẢNG ĐỐI TƯỢNG) ---
        p_lg = self.doc.add_paragraph()
        p_lg.add_run("Lời giải").bold = True
       
        # 6. Đáp án bit (VD: 1001)
        p_da = self.doc.add_paragraph()
        p_da.add_run(str(cau.get("dap_an_dung_sai", ""))).bold = True
        self.doc.add_paragraph("####")
 
        # 7. Render Giải thích Tiếng Việt (MẢNG ĐỐI TƯỢNG)
        giai_thich_arr = cau.get("giai_thich", [])
        if isinstance(giai_thich_arr, list):
            for item in giai_thich_arr:
                # Dòng tiêu đề: + (a.) Nội dung ý. KẾT LUẬN
                p_title = self.doc.add_paragraph()
               
                # Phần prefix: + (a.)
                prefix_run = p_title.add_run(f"+ ({item.get('ky_hieu', '').upper()}.) ")
                prefix_run.bold = False
               
                # Lấy nội dung ý từ cac_y để hiển thị
                ky_hieu_curr = item.get('ky_hieu', '')
                noi_dung_y = ""
                for y in cau.get("cac_y", []):
                    if y.get('ky_hieu') == ky_hieu_curr:
                        noi_dung_y = y.get('noi_dung', '')
                        break
               
                # Thêm nội dung ý
                if noi_dung_y:
                    process_text_with_latex(noi_dung_y, p_title, bold=False)
                    p_title.add_run(" ")
               
                # Thêm kết luận (IN ĐẬM)
                ket_luan_run = p_title.add_run(item.get('ket_luan', ''))
                ket_luan_run.bold = True
               
                # Giải thích chi tiết (dòng tiếp theo)
                p_detail = self.doc.add_paragraph()
                process_text_with_latex(item.get('noi_dung', ''), p_detail, bold=False)
 
        # 8. Render Giải thích Tiếng Anh (nếu có)
        giai_thich_en_arr = cau.get("giai_thich_en", [])
        if isinstance(giai_thich_en_arr, list) and len(giai_thich_en_arr) > 0:
            self.doc.add_paragraph("(translate_en)").italic = True
           
            for item in giai_thich_en_arr:
                # Dòng tiêu đề EN
                p_title_en = self.doc.add_paragraph()
               
                prefix_run_en = p_title_en.add_run(f"+ ({item.get('ky_hieu', '')}.) ")
                prefix_run_en.bold = False
               
                # Lấy nội dung ý EN
                ky_hieu_curr = item.get('ky_hieu', '')
                noi_dung_y_en = ""
                for y in cau.get("cac_y", []):
                    if y.get('ky_hieu') == ky_hieu_curr:
                        noi_dung_y_en = y.get('noi_dung_en', y.get('noi_dung', ''))
                        break
               
                if noi_dung_y_en:
                    process_text_with_latex(noi_dung_y_en, p_title_en, bold=False)
                    p_title_en.add_run(". ")
               
                # Kết luận EN (IN ĐẬM)
                ket_luan_en_run = p_title_en.add_run(item.get('ket_luan', ''))
                ket_luan_en_run.bold = True
               
                # Giải thích chi tiết EN
                p_detail_en = self.doc.add_paragraph()
                process_text_with_latex(item.get('noi_dung', ''), p_detail_en, bold=False)  
    
    def render_question_tra_loi_ngan(self, cau: Dict):
        self.render_ma_dang_header(cau.get("ma_dang"))
        p = self.doc.add_paragraph()
        p.add_run(f"Câu {cau['stt']}. ").bold = True
        process_text_with_latex(cau.get('noi_dung', ''), p)  
        
        # 2. Câu hỏi Tiếng Anh
        if cau.get('noi_dung_en'):
            self.doc.add_paragraph("(translate_en)").italic = True
            p_en = self.doc.add_paragraph()
            process_text_with_latex(cau.get('noi_dung_en', ''), p_en)

        # 3. Hình ảnh
        hinh_anh = cau.get("hinh_anh", {})
        if hinh_anh.get("co_hinh"):
            insert_image_or_placeholder(self.doc, hinh_anh)
        
        # 4. Đáp án
        p_da = self.doc.add_paragraph()
        run_label = p_da.add_run("Đáp án: ")
        run_label.bold = True
        
        raw_ans = str(cau.get('dap_an', '')).strip()
        if not (raw_ans.startswith("[[") and raw_ans.endswith("]]")):
            final_ans = f"[[{raw_ans}]]"
        else:
            final_ans = raw_ans
        process_text_with_latex(final_ans, p_da, bold=True)  
        
        # 5. Lời giải Header
        p_lg = self.doc.add_paragraph()
        p_lg.add_run("Lời giải").bold = True
        self.doc.add_paragraph("####")
        
        # --- HÀM HELPER MỚI: CHUẨN HÓA DÒNG TRỐNG ---
        def render_explanation_block(text_content, lang='vi'):
            if not text_content: return
            
            # 1. Chuẩn hóa xuống dòng: Biến \n\n, \n\s*\n thành 1 \n duy nhất
            # Xử lý ký tự đặc biệt \\n trước
            clean_text = text_content.replace('\\n', '\n')
            # Regex gộp nhiều dòng trống thành 1
            clean_text = re.sub(r'\n\s*\n', '\n', clean_text)
            
            lines = clean_text.split('\n')
            
            for line in lines:
                text = line.strip()
                # Bỏ qua dòng rỗng hoặc dòng chỉ có dấu ####
                if not text or text == "####": continue
                
                is_bold = False
                # Xử lý markdown bold thủ công từ AI
                if text.startswith("**") and text.endswith("**"):
                    text = text[2:-2]
                    is_bold = True
                
                # Auto-detect dòng kết luận để in đậm
                text_lower = text.lower()
                if lang == 'vi' and (text_lower.startswith("vậy") or text_lower.startswith("kết luận") or "(kl.)" in text_lower):
                    is_bold = True
                elif lang == 'en' and (text_lower.startswith("therefore") or text_lower.startswith("conclusion")):
                    is_bold = True
                
                text = text.replace('**', '') 
                p_gt = self.doc.add_paragraph()
                process_text_with_latex(text, p_gt, bold=is_bold)

        # 6. Render Giải thích
        render_explanation_block(cau.get("giai_thich", ""), lang='vi')

        if cau.get("giai_thich_en"):
            self.doc.add_paragraph("(translate_en)").italic = True
            render_explanation_block(cau.get("giai_thich_en", ""), lang='en')  

    def render_question_tu_luan(self, cau: Dict):
        self.render_ma_dang_header(cau.get("ma_dang"))
        """Render câu hỏi Tự luận (Đã fix lỗi khoảng cách dòng)"""
        # 1. Câu hỏi Tiếng Việt
        p = self.doc.add_paragraph()
        p.add_run(f"Câu {cau['stt']}. ").bold = True
        process_text_with_latex(cau.get('noi_dung', ''), p)
        
        # 2. Câu hỏi Tiếng Anh
        if cau.get('noi_dung_en'):
            self.doc.add_paragraph("(translate_en)").italic = True
            p_en = self.doc.add_paragraph()
            process_text_with_latex(cau.get('noi_dung_en', ''), p_en)
        
        # 3. Hình ảnh
        hinh_anh = cau.get("hinh_anh", {})
        if hinh_anh.get("co_hinh"):
            insert_image_or_placeholder(self.doc, hinh_anh)
            
        # 4. Header Lời giải
        p_lg = self.doc.add_paragraph()
        p_lg.add_run("Lời giải").bold = True
        # self.doc.add_paragraph("####")
        
        # --- HÀM HELPER MỚI: CHUẨN HÓA DÒNG TRỐNG ---
        def render_essay_solution(text_content, lang='vi'):
            if not text_content: return
            
            # 1. Chuẩn hóa: Gộp dòng trống thừa
            clean_text = text_content.replace('\\n', '\n')
            clean_text = re.sub(r'\n\s*\n', '\n', clean_text)
            
            lines = clean_text.split('\n')
            
            for line in lines:
                text = line.strip()
                # QUAN TRỌNG: Kích hoạt lại bộ lọc dòng trống
                if not text or text == "####": continue
                
                # Logic in đậm
                is_bold = False
                text_lower = text.lower()
                
                if (text.startswith("**") and text.endswith("**")):
                    text = text[2:-2]
                    is_bold = True
                elif lang == 'vi' and (text_lower.startswith("vậy") or "(kl.)" in text_lower):
                    is_bold = True
                elif lang == 'en' and text_lower.startswith("therefore"):
                    is_bold = True
                
                text = text.replace('**', '')
                p_gt = self.doc.add_paragraph()
                process_text_with_latex(text, p_gt, bold=is_bold)

        # 5. Render Giải thích
        render_essay_solution(cau.get("giai_thich", ""), lang='vi')

        if cau.get("giai_thich_en"):
            self.doc.add_paragraph("(translate_en)").italic = True
            render_essay_solution(cau.get("giai_thich_en", ""), lang='en')   
    
    def render_all(self, data: Dict):
        """
        Main render function - Có hỗ trợ chia PHẦN (PART) bên trong Mức độ
        """
        self.render_title(data)
        
        # 1. Auto-group theo mức độ (Nhận biết, Thông hiểu...)
        grouped = self.auto_group_questions(data)
        
        # 2. Detect loại đề
        loai_de = data.get("loai_de", "")
        ma_bai = data.get("ma_bai", "")
        
        # 3. Render từng nhóm MỨC ĐỘ
        # Thứ tự ưu tiên render
        order_muc_do = ["nhan_biet", "thong_hieu", "van_dung", "van_dung_cao"]
        
        for muc_do in order_muc_do:
            if muc_do not in grouped:
                continue
            
            # Lấy danh sách câu hỏi trong mức độ này
            questions = grouped[muc_do]
            if not questions:
                continue
            section_title = self.get_section_title(muc_do)
            self.doc.add_heading(section_title, level=2)
            current_phan = None

            for cau in questions:
                # Lấy tên phần của câu hiện tại
                raw_phan = cau.get("phan", [])
                
                # Xử lý: Nếu là List thì nối lại thành chuỗi để hiển thị đẹp
                if isinstance(raw_phan, list):
                    # Ví dụ: "Bài 1 - Phần 2 - Dạng bài..."
                    phan_cua_cau = " - ".join([str(x) for x in raw_phan if x])
                else:
                    # Fallback nếu AI lỡ trả về string cũ
                    phan_cua_cau = str(raw_phan).strip()
                
                # Nếu câu này thuộc một phần mới -> In Header Phần
                if phan_cua_cau and phan_cua_cau != current_phan:
                    # In ra header cấp 3 (VD: Phần 1: Đội ngũ...)
                    # Dùng màu hoặc in đậm để phân biệt
                    p_phan = self.doc.add_heading(phan_cua_cau.upper(), level=3)
                    p_phan.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    current_phan = phan_cua_cau
                
                # Render nội dung câu hỏi như bình thường
                if loai_de == "dung_sai":
                    self.render_question_dung_sai(cau)
                elif loai_de == "tra_loi_ngan":
                    self.render_question_tra_loi_ngan(cau)
                elif loai_de == "tu_luan":  # [THÊM MỚI]
                    self.render_question_tu_luan(cau)
                else:
                    self.render_question_trac_nghiem(cau)

def clean_json_response(text):
    """Làm sạch chuỗi trả về từ AI, loại bỏ markdown và ký tự thừa"""
    try:
        # Loại bỏ các tag ```json hoặc ``` nếu có
        clean_text = re.sub(r'```json|```', '', text).strip()
        return clean_text
    except Exception:
        return text
 
import re

def roman_to_int(s):
    """Chuyển số La Mã sang số nguyên (Hỗ trợ I..XX)"""
    s = s.upper().strip().replace('.', '').replace(':', '').replace(')', '')
    romans = {'I': 1, 'V': 5, 'X': 10}
    
    # Map nhanh các số nhỏ hay gặp để tốc độ cao nhất
    fast_map = {
        'I': 1, 'II': 2, 'III': 3, 'IV': 4, 'V': 5,
        'VI': 6, 'VII': 7, 'VIII': 8, 'IX': 9, 'X': 10,
        'XI': 11, 'XII': 12
    }
    if s in fast_map: return fast_map[s]
    return None

def renumber_ma_dang_global(all_questions, reference_ma_bai):
    """
    [SPECIALIZED FUNCTION FOR DUNG_SAI ONLY]
    Logic: 
    1. Bỏ qua ID Mục/Phần.
    2. Chỉ đếm ID Dạng tăng dần (Global Counter).
    3. Cấu trúc output: [MA_BAI]_[ID_DẠNG] (Ví dụ: SN_HOA_10_1_1_1, SN_HOA_10_1_1_2...)
    """
    print(f"🔧 [DungSai Only] Chuẩn hóa Mã Dạng (No-Section) cho: {reference_ma_bai}")
    
    # Bộ nhớ map tên dạng -> ID (Dùng chung cho cả bài)
    type_memory = {} 
    
    # Bộ đếm ID dạng toàn cục (bắt đầu từ 0)
    global_dang_counter = 0
    
    final_questions = []
    
    # Biến fallback phòng khi AI trả thiếu mảng phan
    last_known_phan = [reference_ma_bai, "Mục 1", "Dạng tổng quát"]

    for index, q in enumerate(all_questions):
        # 1. Cập nhật STT chuẩn (cho chắc chắn)
        q['stt'] = index + 1
        
        # 2. Lấy dữ liệu phân cấp
        raw_phan = q.get("phan", [])
        if not isinstance(raw_phan, list) or len(raw_phan) < 3:
            raw_phan = list(last_known_phan)
        else:
            last_known_phan = raw_phan

        # Tách các thành phần (Chỉ dùng để hiển thị hoặc map key)
        ten_bai = str(raw_phan[0]).strip()
        ten_muc = str(raw_phan[1]).strip()
        ten_dang = str(raw_phan[2]).strip() # Key quan trọng nhất
        
        # 3. THUẬT TOÁN GÁN ID (GLOBAL - NO SECTION)
        # Chỉ quan tâm tên dạng. Nếu tên dạng trùng -> ID cũ. Nếu mới -> ID mới.
        if ten_dang in type_memory:
            current_dang_id = type_memory[ten_dang]
        else:
            global_dang_counter += 1
            current_dang_id = global_dang_counter
            type_memory[ten_dang] = current_dang_id

        # 4. Tạo chuỗi ma_dang chuẩn: MA_BAI + "_" + ID_DANG
        final_ma_dang = f"{reference_ma_bai}_{current_dang_id}"
        
        q['ma_dang'] = final_ma_dang
        # Cập nhật lại phan để đảm bảo thống nhất
        q['phan'] = [ten_bai, ten_muc, ten_dang] 
        
        final_questions.append(q)

    print(f"   ✅ [DungSai] Đã map {global_dang_counter} dạng bài duy nhất.")
    return final_questions

def process_dung_sai_smart_batch(file_path, base_prompt, file_name, project_id, creds, model_name, batch_name):
    """
    Quy trình xử lý ĐÚNG/SAI (Fixed Logic): 
    - Batch 1 (1-20): 100% THÔNG HIỂU.
    - Batch 2 (21-40): 10 VẬN DỤNG + 10 VẬN DỤNG CAO.
    - Clean trường 'phan' để không bị dính từ khóa mức độ.
    """
    from modules.common.callAPI import VertexClient
    client = VertexClient(project_id, creds, model_name)
    
    # --- CẤU HÌNH BATCH (TUÂN THỦ 20 THÔNG HIỂU) ---
    batches = [
        {
            "range": "1-20", 
            # Instruction gửi cho AI: Ép sinh Thông hiểu
            "level_desc": "YÊU CẦU: Sinh 20 câu THÔNG HIỂU (Giải thích, so sánh, phân biệt bản chất). KHÔNG sinh câu Nhận biết (định nghĩa đơn giản).", 
            "mode": "thong_hieu" 
        },
        {
            "range": "21-40", 
            # Instruction gửi cho AI: Ép sinh Vận dụng & Vận dụng cao
            "level_desc": "YÊU CẦU: 10 câu VẬN DỤNG (Tính toán, áp dụng công thức) và 10 câu VẬN DỤNG CAO (Suy luận, tổng hợp, bài toán ngược).", 
            "mode": "van_dung" # Mode tạm, sẽ override chi tiết bên dưới
        }
    ]

    all_raw_questions = []
    reference_ma_bai = "SN_UNK" 
    
    print(f"\n🚀 [DungSai] Chạy chế độ Smart Batch (20 Thông hiểu - 10 Vận dụng - 10 VDC) cho: {batch_name}")

    for idx, batch in enumerate(batches):
        print(f"   ► Batch {idx+1}: Câu {batch['range']}...")
        
        batch_instruction = f"""
{base_prompt}
--------------------------------------------------------------------------------
⚠️ LỆNH THỰC THI BATCH {idx+1}/2:
1. PHẠM VI STT: {batch['range']}.
2. YÊU CẦU MỨC ĐỘ: {batch['level_desc']}
3. QUY ĐỊNH NGHIÊM NGẶT: 
   - Trường "phan" CHỈ ĐƯỢC CHỨA: ["Tên Bài", "Tên Mục", "Tên Dạng"]. 
   - CẤM đưa từ khóa "Thông hiểu", "Vận dụng" vào trường "phan".
   - Trích xuất Mã Bài (ma_bai) chuẩn SN_[MÔN]...
--------------------------------------------------------------------------------
"""
        try:
            raw_text = client.send_data_to_AI(batch_instruction, file_path, response_schema=schema_dung_sai, max_output_tokens=65534)
            if not raw_text: continue

            data = json.loads(clean_json_response(raw_text))
            batch_questions = data.get("cau_hoi", [])
            
            # --- CODE FIX: CLEAN 'PHAN' & FORCE 'MUC_DO' ---
            keywords_to_remove = ["nhận biết", "thông hiểu", "vận dụng", "mức độ", "level", "nhan_biet", "thong_hieu", "van_dung"]
            
            for q in batch_questions:
                # 1. Clean trường 'phan' (Xóa các từ khóa mức độ nếu AI lỡ điền vào)
                raw_phan = q.get("phan", [])
                if isinstance(raw_phan, list):
                    clean_phan = []
                    for p_item in raw_phan:
                        p_str = str(p_item)
                        # Nếu dòng này chứa từ khóa cấm -> Không add vào, hoặc clean nhẹ
                        # Ở đây ta chọn giải pháp: Nếu chứa từ khóa mức độ -> Bỏ qua phần tử đó luôn
                        if not any(kw in p_str.lower() for kw in keywords_to_remove):
                             clean_phan.append(p_str)
                    
                    # Nếu clean xong bị rỗng hoặc mất dạng -> Fallback nhẹ
                    if len(clean_phan) < 3:
                        # Giữ nguyên bản gốc nếu clean làm hỏng cấu trúc (chấp nhận xấu còn hơn lỗi code)
                        pass 
                    else:
                        q['phan'] = clean_phan
                
                # 2. Force Mức độ (Ghi đè cứng theo logic STT của bạn)
                stt = q.get("stt", 0)
                
                # Logic: 20 Thông hiểu -> 10 Vận dụng -> 10 VDC
                if 1 <= stt <= 20:
                    q['muc_do'] = "thong_hieu"
                elif 21 <= stt <= 30:
                    q['muc_do'] = "van_dung"
                elif 31 <= stt <= 40:
                    q['muc_do'] = "van_dung_cao"
                # (Các câu ngoài range này sẽ giữ nguyên giá trị AI trả về)

            if idx == 0:
                raw_ma = data.get("ma_bai", "")
                if raw_ma and raw_ma.startswith("SN_") and raw_ma.count("_") >= 3:
                    reference_ma_bai = raw_ma
                else:
                    if reference_ma_bai == "SN_UNK": reference_ma_bai = f"SN_UNK_{batch_name}"

            all_raw_questions.extend(batch_questions)
            print(f"      ✅ Batch {idx+1} OK: +{len(batch_questions)} câu.")
            
        except Exception as e:
            print(f"      ❌ Lỗi Batch {idx+1}: {e}")

    if not all_raw_questions: return None

    # Gọi Renumber (Logic bỏ Mục, chỉ giữ ID dạng tăng dần)
    final_questions = renumber_ma_dang_global(all_raw_questions, reference_ma_bai)
    
    return {
        "loai_de": "dung_sai",
        "tong_so_cau": len(final_questions),
        "ma_bai": reference_ma_bai,
        "cau_hoi": final_questions
    }

# ==============================================================================
# MAIN ENTRY POINT
# ==============================================================================

def response2docx_flexible(file_path, prompt, file_name, project_id, creds, model_name, question_type="trac_nghiem_4_dap_an", batch_name=None):
    if not batch_name:
        batch_name = file_name.replace("_TN", "").replace("_DS", "").replace("_TLN", "")
        
    try:
        final_json_data = None

        # 1. LOGIC RIÊNG CHO ĐÚNG/SAI (Có can thiệp code renumber)
        if question_type == "dung_sai":
            final_json_data = process_dung_sai_smart_batch(
                file_path, prompt, file_name, project_id, creds, model_name, batch_name
            )

        # 2. LOGIC CHO CÁC DẠNG KHÁC (Tuyệt đối tin tưởng Prompt AI, không renumber)
        else:
            from modules.common.callAPI import VertexClient
            client = VertexClient(project_id, creds, model_name)
            target_schema = get_schema_by_type(question_type)
            final_prompt = PromptBuilder.wrap_user_prompt(prompt)
            
            print(f"📤 [{question_type}] Đang gửi request (1-shot)...")
            ai_response_text = client.send_data_to_AI(
                final_prompt, file_path, response_schema=target_schema, max_output_tokens=65534
            )
            
            if ai_response_text:
                final_json_data = json.loads(clean_json_response(ai_response_text))
            
            # KHÔNG GỌI renumber_ma_dang_global ở đây.
            # Dữ liệu AI trả về sao thì dùng vậy.

        # --- PHẦN CHUNG: LƯU FILE ---
        if not final_json_data: 
            print("❌ Không có dữ liệu.")
            return None
        
        print(f"💾 [{batch_name}] Lưu JSON...")
        save_json_securely(final_json_data, batch_name, file_name)
        
        print(f"📝 [{batch_name}] Render DOCX...")
        doc = Document()
        renderer = DynamicDocxRenderer(doc)
        renderer.render_all(final_json_data)
        
        output_path = save_document_securely(doc, batch_name, file_name)
        if output_path: print(f"✅ HOÀN THÀNH: {output_path}")
        return output_path

    except Exception as e:
        print(f"❌ Lỗi hệ thống: {e}")
        traceback.print_exc()
        return None

# def response2docx_flexible(
#     file_path: str,
#     prompt: str,
#     file_name: str,
#     project_id: str,
#     creds: str,
#     model_name: str,
#     question_type: str = "trac_nghiem_4_dap_an",
#     batch_name: str = None
# ):
#     try:
#         from modules.common.callAPI import VertexClient
#         client = VertexClient(project_id, creds, model_name)
#         if not batch_name:
#             batch_name = file_name.replace("_TN", "").replace("_DS", "").replace("_TLN", "")
            
#         target_schema = get_schema_by_type(question_type)
#         final_prompt = PromptBuilder.wrap_user_prompt(prompt)
        
#         print(f"📤 Đang gửi request (Schema: {question_type})...")
#         ai_response_text = client.send_data_to_AI(
#             final_prompt, 
#             file_path, 
#             response_schema=target_schema,
#             max_output_tokens=65534
#         )
        
#         if not ai_response_text:
#             print("❌ AI không trả về dữ liệu.")
#             return None

#         print("🔄 Đang parse JSON...")
#         try:
#             data = json.loads(ai_response_text)
#         except json.JSONDecodeError as e:
#             print(f"❌ Lỗi JSON bất ngờ: {e}")
#             if "```json" in ai_response_text:
#                 clean_text = ai_response_text.split("```json")[1].split("```")[0]
#                 data = json.loads(clean_text)
#             else:
#                 return None

#         print(f"✅ Parse thành công: {data.get('tong_so_cau', 0)} câu hỏi")
        
#         # --- [MỚI] LƯU FILE JSON ---
#         print("💾 Đang lưu file JSON...")
#         save_json_securely(data, batch_name, file_name)
#         # ---------------------------

#         print("📝 Đang tạo DOCX...")
#         doc = Document()
#         renderer = DynamicDocxRenderer(doc)
        
#         try:
#             renderer.render_all(data)
#         except Exception as render_err:
#              print(f"❌ Lỗi Render: {render_err}")
#              traceback.print_exc()

#         print("💾 Đang lưu file DOCX...")
#         output_path = save_document_securely(doc, batch_name, file_name)
#         return output_path

#     except Exception as e:
#         print(f"❌ Lỗi hệ thống: {e}")
#         traceback.print_exc()
#         return None

#     except Exception as e_main:
#         print(f"❌ LỖI NGHIÊM TRỌNG TRONG TOÀN BỘ HÀM: {e_main}")
#         traceback.print_exc()
#         try:
#             doc = Document()
#             doc.add_heading('LỖI HỆ THỐNG', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
#             doc.add_paragraph(f'Lỗi nghiêm trọng: {e_main}')
#             doc.add_paragraph('Hệ thống không thể xử lý yêu cầu.')
            
#             if not batch_name:
#                 batch_name = file_name.replace("_TN", "").replace("_DS", "").replace("_TLN", "")
#             fallback_path = os.path.join(ensure_output_folder_for_batch(batch_name), f"{file_name}_loi_he_thong.docx")
            
#             doc.save(fallback_path)
#             return fallback_path
#         except Exception as e_final:
#             return None

def response2docx_json(file_path, prompt, file_name, project_id, creds, model_name, batch_name=None):
    """Wrapper cho trắc nghiệm 4 đáp án (legacy)"""
    return response2docx_flexible(
        file_path, prompt, file_name, project_id, creds, model_name,
        question_type="trac_nghiem_4_dap_an",
        batch_name=batch_name
    )

def response2docx_dung_sai_json(file_path, prompt, file_name, project_id, creds, model_name, batch_name=None):
    """Wrapper cho đúng/sai (legacy)"""
    return response2docx_flexible(
        file_path, prompt, file_name, project_id, creds, model_name,
        question_type="dung_sai",
        batch_name=batch_name
    )
    
def response2docx_tra_loi_ngan_json(file_path, prompt, file_name, project_id, creds, model_name, batch_name=None):
    """Wrapper cho trả lời ngắn (legacy compatibility)"""
    return response2docx_flexible(
        file_path, prompt, file_name, project_id, creds, model_name,
        question_type="tra_loi_ngan",
        batch_name=batch_name
    )

def response2docx_tu_luan_json(file_path, prompt, file_name, project_id, creds, model_name, batch_name=None):
    """Wrapper cho tự luận học liệu"""
    return response2docx_flexible(
        file_path, prompt, file_name, project_id, creds, model_name,
        question_type="tu_luan", # Key này sẽ kích hoạt logic trong PromptBuilder và Renderer
        batch_name=batch_name
    )