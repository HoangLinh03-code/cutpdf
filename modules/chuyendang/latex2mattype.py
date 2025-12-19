import subprocess
import tempfile
import os
from docx import Document

def _convert_latex_string_to_mathtype_docx(latex_text_content, output_file_path):
    """
    Hàm trợ giúp để chuyển đổi một chuỗi chứa LaTeX sang tệp DOCX
    với các phương trình MathType bằng Pandoc.

    Args:
        latex_text_content (str): Chuỗi đầu vào chứa văn bản và các biểu thức LaTeX.
        output_file_path (str): Đường dẫn cho tệp DOCX đầu ra.

    Returns:
        bool: True nếu chuyển đổi thành công, False nếu ngược lại.
    """
    try:
        # Tạo một tệp markdown tạm thời với nội dung LaTeX
        # Sử dụng encoding='utf-8' để đảm bảo xử lý ký tự đúng cách
        with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as temp_md:
            temp_md.write(latex_text_content)
            temp_md_path = temp_md.name

        print(f"🔄 Đang cố gắng chuyển đổi LaTeX sang MathType DOCX bằng Pandoc...")
        # Sử dụng pandoc để chuyển đổi markdown+latex sang docx
        # --from=markdown+tex_math_dollars cho Pandoc biết cách diễn giải $...$ và $$...$$ là toán học LaTeX
        result = subprocess.run(
            [
                'pandoc',
                temp_md_path,
                '-o', output_file_path,
                '--from=markdown+tex_math_dollars',
                '--to=docx'
            ],
            capture_output=True,
            text=True,
            encoding='utf-8' # Đảm bảo đầu ra được giải mã chính xác
        )

        # Xóa tệp tạm thời
        os.unlink(temp_md_path)

        if result.returncode == 0:
            print(f"✅ Đã tạo tệp thành công: {output_file_path}")
            return True
        else:
            print(f"❌ Lỗi Pandoc trong quá trình chuyển đổi: {result.stderr}")
            return False
    except FileNotFoundError:
        print("❌ Lỗi: Không tìm thấy Pandoc. Vui lòng đảm bảo Pandoc đã được cài đặt và có thể truy cập trong PATH của hệ thống bạn.")
        print("   Tải xuống Pandoc từ: https://pandoc.org/installing.html")
        return False
    except Exception as e:
        print(f"❌ Đã xảy ra lỗi không mong muốn: {e}")
        return False

def convert_docx_with_latex_to_mathtype_docx(input_docx_path, output_docx_path):
    """
    Đọc một tệp DOCX đầu vào, trích xuất tất cả nội dung văn bản (bao gồm LaTeX nhúng),
    và chuyển đổi nó thành một tệp DOCX mới trong đó các biểu thức LaTeX được hiển thị
    dưới dạng các phương trình MathType bằng Pandoc.

    Args:
        input_docx_path (str): Đường dẫn đến tệp DOCX đầu vào chứa LaTeX.
        output_docx_path (str): Đường dẫn cho tệp DOCX đầu ra mới.

    Returns:
        bool: True nếu chuyển đổi thành công, False nếu ngược lại.
    """
    try:
        print(f"📖 Đang đọc nội dung từ DOCX đầu vào: {input_docx_path}")
        doc = Document(input_docx_path)
        full_text_content = []

        # Trích xuất văn bản từ các đoạn văn
        for para in doc.paragraphs:
            full_text_content.append(para.text)

        # Trích xuất văn bản từ các bảng (nếu có)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text_content.append(cell.text)

        # Nối tất cả văn bản đã trích xuất thành một chuỗi duy nhất
        extracted_text = "\n".join(full_text_content)

        if not extracted_text.strip():
            print("⚠️ Cảnh báo: Không tìm thấy nội dung văn bản trong tệp DOCX đầu vào.")
            return False

        print("✨ Đã trích xuất nội dung văn bản. Đang tiến hành chuyển đổi LaTeX.")
        # Sử dụng hàm trợ giúp để thực hiện chuyển đổi
        return _convert_latex_string_to_mathtype_docx(extracted_text, output_docx_path)

    except FileNotFoundError:
        print(f"❌ Lỗi: Không tìm thấy tệp DOCX đầu vào tại '{input_docx_path}'. Vui lòng kiểm tra đường dẫn.")
        return False
    except Exception as e:
        print(f"❌ Đã xảy ra lỗi khi xử lý tệp DOCX: {e}")
        return False

if __name__ == "__main__":
    # Thay đổi các đường dẫn này để trỏ đến tệp DOCX của bạn
    input_your_doc_path = r"C:\Users\Admin\Downloads\Xử lý chuyển xml\input_docx\Toán lớp 4 tập 1 - Chân trời sáng tạo_done.docx" # Đổi thành đường dẫn tệp DOCX của bạn
    output_doc_path = "output_math_converted.docx" # Tên tệp DOCX đầu ra

    print("--- Bắt đầu chuyển đổi DOCX sang MathType ---")
    success = convert_docx_with_latex_to_mathtype_docx(input_your_doc_path, output_doc_path)

    if success:
        print(f"\nChuyển đổi hoàn tất. Kiểm tra '{output_doc_path}' để xem kết quả.")
        print("Lưu ý: Tệp DOCX mới sẽ chứa văn bản đã trích xuất với LaTeX được chuyển đổi thành MathType.")
        print("Định dạng gốc, hình ảnh và bố cục phức tạp từ DOCX đầu vào có thể không được bảo toàn hoàn hảo.")
    else:
        print("\nChuyển đổi thất bại. Vui lòng xem lại các thông báo lỗi ở trên.")
