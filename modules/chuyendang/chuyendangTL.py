import docx
import re
import os
import json
import time

from datetime import datetime

from docx import Document
from docx.shared import Inches
from docx.oxml.shared import qn
from docx.oxml import OxmlElement
from PIL import Image
import config
from modules.chuyendang.md2docx import md2docx
from modules.chuyendang.docx2md import convert_docx_to_md

from modules.chuyendang.ai_handle.vertex_ai import key
from google.oauth2.credentials import Credentials
import vertexai
from vertexai.generative_models import GenerativeModel, GenerationConfig, Part
import io
# Các hàm thành phần khác không thay đổi
global_image_counter = 0

def get_unique_image_filename(temp_dir, ext):
    global global_image_counter
    filename = f"{temp_dir}/img_{global_image_counter}.{ext}"
    global_image_counter += 1
    return filename

def load_api_keys(json_path):
    with open(json_path, 'r', encoding='utf-8') as file:
        data = json.load(file)
        return data.get('keys', [])

def read_prompt(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()
import requests

def get_timestamp():
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")

def extract_images_from_paragraph(paragraph, temp_dir="temp_images"):
    if not os.path.exists(temp_dir):
        os.makedirs(temp_dir)
    
    images = []
    try:
        for run in paragraph.runs:
            drawing_elements = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
            for drawing in drawing_elements:
                blips = drawing.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                for blip in blips:
                    embed_attr = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    if embed_attr:
                        try:
                            image_part = paragraph.part.related_parts[embed_attr]
                            content_type = image_part.content_type
                            ext = 'png' if 'png' in content_type else 'jpg' if 'jpeg' in content_type or 'jpg' in content_type else 'gif' if 'gif' in content_type else 'png'
                            image_filename = get_unique_image_filename(temp_dir, ext)
                            with open(image_filename, 'wb') as img_file:
                                img_file.write(image_part.blob)
                            images.append(image_filename)
                            print(f"[DEBUG] Extracted image: {image_filename}")
                        except Exception as e:
                            print(f"Lỗi khi trích xuất ảnh từ relationship {embed_attr}: {e}")
    except Exception as e:
        print(f"Lỗi khi trích xuất ảnh từ paragraph: {e}")
    return images

def table_to_markdown(table):
    try:
        markdown_rows = []
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip().replace('\n', ' ').replace('|', '\\|') for cell in row.cells]
            markdown_rows.append('| ' + ' | '.join(cells) + ' |')
            if i == 0:
                separator = '| ' + ' | '.join(['---'] * len(cells)) + ' |'
                markdown_rows.append(separator)
        return '\n'.join(markdown_rows)
    except Exception as e:
        print(f"Lỗi khi chuyển bảng thành markdown: {e}")
        return ""


def extract_question_content_with_media(doc, start_index, end_index, temp_dir="temp_images"):
    content_parts = []
    extracted_images = []
    try:
        body = doc.element.body
        all_elements = list(body)
        
        start_element_pos = -1
        for i, element in enumerate(all_elements):
            if element == doc.paragraphs[start_index]._element:
                start_element_pos = i
                break
        
        end_element_pos = -1
        if end_index < len(doc.paragraphs):
            for i, element in enumerate(all_elements):
                if element == doc.paragraphs[end_index]._element:
                    end_element_pos = i
                    break
        else:
            end_element_pos = len(all_elements)

        if start_element_pos == -1 or end_element_pos == -1:
            print(f"[ERROR] Could not find start/end paragraphs in document structure.")
            return "", []

        for i in range(start_element_pos, end_element_pos):
            element = all_elements[i]
            
            if element.tag.endswith('}p'):
                # Find the paragraph corresponding to the element
                para = None
                for p in doc.paragraphs:
                    if p._element == element:
                        para = p
                        break
                if para:
                    para_text = para.text.strip()
                    images = extract_images_from_paragraph(para, temp_dir)
                    for img_path in images:
                        img_idx_for_placeholder = int(img_path.split('_')[-1].split('.')[0])
                        para_text += f" [IMG-{img_idx_for_placeholder}]"
                        extracted_images.append(img_path)
                    if para_text:
                        content_parts.append(para_text)
            
            elif element.tag.endswith('}tbl'):
                # Find the table corresponding to the element
                table = None
                for t in doc.tables:
                    if t._tbl == element:
                        table = t
                        break
                if table:
                    markdown_table = table_to_markdown(table)
                    if markdown_table:
                        content_parts.append(f"\n{markdown_table}\n")

    except Exception as e:
        print(f"Lỗi khi xử lý nội dung trong phạm vi: {e}")
    
    full_content = '\n'.join(content_parts)
    return full_content, extracted_images

def copy_paragraph_with_media(source_para, target_doc, source_doc):
    new_para = target_doc.add_paragraph()
    new_para.style = source_para.style
    
    for run in source_para.runs:
        new_run = new_para.add_run(run.text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        if run.font.size:
            new_run.font.size = run.font.size
    return new_para

def copy_content_in_range(source_doc, target_doc, start_index, end_index):
    for i in range(start_index, end_index):
        if i < len(source_doc.paragraphs):
            para = source_doc.paragraphs[i]
            if para._element.getparent().tag.endswith('}tbl'):
                for table in source_doc.tables:
                    if table._tbl == para._element.getparent():
                        new_table = target_doc.add_table(rows=len(table.rows), cols=len(table.columns))
                        new_table.style = table.style
                        for row_idx, row in enumerate(table.rows):
                            for cell_idx, cell in enumerate(row.cells):
                                new_table.cell(row_idx, cell_idx).text = cell.text
                        break
            else:
                copy_paragraph_with_media(para, target_doc, source_doc)

def call_gemini_api_with_images(prompt, subject, level, question, images=None, last_used_key_index=None):
    """
    Gọi Gemini trên Vertex AI với prompt và ảnh (nếu có), tự động chọn key từ danh sách api_keys.
    """
   
    # Chọn key tiếp theo (xoay vòng)
   
    model_name = "gemini-2.5-pro"
    project_id = os.getenv("PROJECT_ID")
    region = os.getenv("REGION") or "us-central1"
    role = "chuyên gia giáo dục"  # Có thể thay đổi nếu cần
    if not project_id:
        print("Lỗi: Không tìm thấy Project ID trong biến môi trường.\n")
    if not region:
        region = "us-central1"
        print("Cảnh báo: Không tìm thấy Region, đang sử dụng mặc định: us-central1\n")

    try:
        credentials = Credentials(token=key)
        vertexai.init(project=project_id, location=region, credentials=credentials)
        model = GenerativeModel(model_name)
        generation_config = GenerationConfig(
            temperature=0.5,
            top_p=0.8,
            response_mime_type="text/plain",
        )

        # Chuẩn bị nội dung prompt
        processed_prompt_text = ""
        if role:
            processed_prompt_text += f"Với vai trò là một {role}, hãy phân tích nội dung được cung cấp."
        formatted_prompt = prompt.format(subject=subject, level=level, question=question)
        processed_prompt_text += formatted_prompt

        contents = [processed_prompt_text]

        # Xử lý hình ảnh nếu có
        allowed_image_formats = ['png', 'jpeg', 'jpg', 'gif', 'webp']
        if images:
            for image_path in images:
                if os.path.exists(image_path):
                    try:
                        img = Image.open(image_path)
                        format = img.format.lower() if img.format else None
                        if format not in allowed_image_formats:
                            img_output = io.BytesIO()
                            img.save(img_output, format='WEBP')
                            img_output.seek(0)
                            img_bytes = img_output.getvalue()
                            mime_type = 'image/webp'
                        else:
                            with open(image_path, "rb") as image_file:
                                img_bytes = image_file.read()
                            mime_type = f'image/{format}'
                        image_part = Part.from_data(img_bytes, mime_type=mime_type)
                        contents.append(image_part)
                    except Exception as image_err:
                        print(f"Lỗi khi xử lý hình ảnh '{image_path}': {image_err}\n")
                else:
                    print(f"Cảnh báo: Không tìm thấy file hình ảnh tại '{image_path}'\n")
        # print(contents)
        print(f"Đang gọi Vertex AI Gemini model '{model_name}' với prompt: ...\n")
        response = model.generate_content(contents, generation_config=generation_config, stream=False)
       
        if response and response.text:
            return response.text.strip(),last_used_key_index
        else:
            print(f"Cảnh báo (Vertex AI): Không có phản hồi văn bản từ mô hình '{model_name}'\n")
            return None,last_used_key_index

    except Exception as e:
        print(f"Lỗi khi gọi API Vertex AI Gemini cho mô hình '{model_name}': {e}\n")
        return None,last_used_key_index

def get_form_code_with_retries(prompt_chon_dang, subject, level, content, images, last_key_index, max_retries=3):
    """
    Hàm riêng để gọi API lấy dạng câu hỏi và thử lại nếu thất bại.
    Trả về form_code và chỉ mục key cuối cùng được sử dụng.
    """
    form_code = None
    retries = 0
    while form_code is None and retries < max_retries:
        print(f"[{get_timestamp()}] [INFO] Thử lại lần {retries + 1} để lấy dạng câu hỏi...")
        time.sleep(2)
        api_response, last_key_index = call_gemini_api_with_images(
            prompt_chon_dang, subject, level, content, 
             images, last_key_index
        )
        form_code = extract_form_code(api_response)
        print(f"[{get_timestamp()}] [DEBUG] Kết quả API (dạng câu hỏi) lần thử {retries + 1}: '{form_code}'")
        
        # Kiểm tra nếu lấy được form_code hợp lệ thì thoát vòng lặp
        if form_code and form_code != 'DD':
            break
        retries += 1
    
    # Trả về form_code và chỉ mục key cuối cùng được sử dụng
    return form_code, last_key_index


def process_api_response_with_media(response_text, output_doc, images_list, images_folder="temp_images"):
    lines = response_text.split('\n')
    print("ket qua ai tra ve " + response_text)
    
    current_paragraph = None

    i = 0
    while i < len(lines):
        line = lines[i]
        
        if '|' in line and (i == 0 or not lines[i-1].strip()):
            table_lines = []
            j = i
            while j < len(lines) and ('|' in lines[j] or not lines[j].strip()):
                if '|' in lines[j]:
                    table_lines.append(lines[j])
                j += 1
            
            if table_lines:
                if current_paragraph is not None and current_paragraph.text.strip():
                    pass
                markdown_table = '\n'.join(table_lines)
                table = markdown_to_table(output_doc, markdown_table)
                current_paragraph = None
                i = j
                continue
        
        if current_paragraph is None:
            current_paragraph = output_doc.add_paragraph()
        images_list = []
        folder_path = images_folder
        if os.path.isdir(folder_path):
            for item in os.listdir(folder_path):
                item_path = os.path.join(folder_path, item)
                if os.path.isfile(item_path):
                    images_list.append(item_path)
        
        parts = re.split(r'(\$\$[^$]*?\$\$|\\\[.*?\\\]|\$[^$]*?\$|\[IMG-\d+\])', line)
        for part in parts:
            if not part:
                continue

            img_match = re.match(r'\[IMG-(\d+)\]', part)
            
            if img_match:
                img_number_from_placeholder = int(img_match.group(1))
                found_img_path = None
                for img_path in images_list:
                    try:
                        img_saved_number = int(os.path.basename(img_path).split('_')[-1].split('.')[0])
                        if img_saved_number == img_number_from_placeholder:
                            found_img_path = img_path
                            break
                    except ValueError:
                        continue 
                
                if found_img_path:
                    try:
                        print(f"[DEBUG] Adding image {found_img_path} to Word document")
                        run = current_paragraph.add_run()
                        run.add_picture(found_img_path, width=Inches(3)) 
                    except Exception as e:
                        print(f"Lỗi khi thêm ảnh {found_img_path} vào tài liệu Word: {e}")
                        current_paragraph.add_run(part) 
                else:
                    current_paragraph.add_run(part) 
            else:
                current_paragraph.add_run(part)
        
        if line.strip(): 
            current_paragraph = None 
        i += 1

def extract_form_code(api_response):
    api_response = str(api_response)
    try:
        data = json.loads(api_response)
        return data.get("Dạng mới")
    except Exception:
        # Fallback cho trường hợp không phải JSON
        match = re.search(r'"Dạng mới"\s*:\s*"(\w+)"', api_response)
        if match:
            return match.group(1)
        match = re.search(r'Dạng mới:\s*(\w+)', api_response)
        return match.group(1) if match else None

def has_valid_solution(text):
    lines = text.split('\n')
    for i, line in enumerate(lines):
        if 'Lời giải' in line.strip():
            if i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                if next_line.isdigit():
                    if i + 2 < len(lines):
                        third_line = lines[i + 2].strip()
                        if '####' in third_line:
                            return True
    return False

def cleanup_temp_files(temp_dir="temp_images"):
    try:
        if os.path.exists(temp_dir):
            for file in os.listdir(temp_dir):
                os.remove(os.path.join(temp_dir, file))
            os.rmdir(temp_dir)
    except Exception as e:
        print(f"Lỗi khi xóa file tạm: {e}")

# Hàm mới để tải và xử lý cấu trúc JSON
def load_json_structure(json_path):
    with open(json_path, 'r', encoding='utf-8') as f:
        return json.load(f)
def remove_vietnamese_diacritics(text):
        """
        Replaces Vietnamese characters with their non-diacritic counterparts.
        """
        replacements = {
        'à': 'a', 'á': 'a', 'ả': 'a', 'ã': 'a', 'ạ': 'a',
        'ă': 'a', 'ằ': 'a', 'ắ': 'a', 'ẳ': 'a', 'ẵ': 'a', 'ặ': 'a',
        'â': 'a', 'ầ': 'a', 'ấ': 'a', 'ẩ': 'a', 'ẫ': 'a', 'ậ': 'a',
        'đ': 'd',
        'è': 'e', 'é': 'e', 'ẻ': 'e', 'ẽ': 'e', 'ẹ': 'e',
        'ê': 'e', 'ề': 'e', 'ế': 'e', 'ể': 'e', 'ễ': 'e', 'ệ': 'e',
        'ì': 'i', 'í': 'i', 'ỉ': 'i', 'ĩ': 'i', 'ị': 'i',
        'ò': 'o', 'ó': 'o', 'ỏ': 'o', 'õ': 'o', 'ọ': 'o',
        'ô': 'o', 'ồ': 'o', 'ố': 'o', 'ổ': 'o', 'ỗ': 'o', 'ộ': 'o',
        'ơ': 'o', 'ờ': 'o', 'ớ': 'o', 'ở': 'o', 'ỡ': 'o', 'ợ': 'o',
        'ù': 'u', 'ú': 'u', 'ủ': 'u', 'ũ': 'u', 'ụ': 'u',
        'ư': 'u', 'ừ': 'u', 'ứ': 'u', 'ử': 'u', 'ữ': 'u', 'ự': 'u',
        'ỳ': 'y', 'ý': 'y', 'ỷ': 'y', 'ỹ': 'y', 'ỵ': 'y',
        'À': 'A', 'Á': 'A', 'Ả': 'A', 'Ã': 'A', 'Ạ': 'A',
        'Ă': 'A', 'Ằ': 'A', 'Ắ': 'A', 'Ẳ': 'A', 'Ẵ': 'A', 'Ặ': 'A',
        'Â': 'A', 'Ầ': 'A', 'Ấ': 'A', 'Ẩ': 'A', 'Ẫ': 'A', 'Ậ': 'A',
        'Đ': 'D',
        'È': 'E', 'É': 'E', 'Ẻ': 'E', 'Ẽ': 'E', 'Ẹ': 'E',
        'Ê': 'E', 'Ề': 'E', 'Ế': 'E', 'Ể': 'E', 'Ễ': 'E', 'Ệ': 'E',
        'Ì': 'I', 'Í': 'I', 'Ỉ': 'I', 'Ĩ': 'I', 'Ị': 'I',
        'Ò': 'O', 'Ó': 'O', 'Ỏ': 'O', 'Õ': 'O', 'Ọ': 'O',
        'Ô': 'O', 'Ồ': 'O', 'Ố': 'O', 'Ổ': 'O', 'Ỗ': 'O', 'Ộ': 'O',
        'Ơ': 'O', 'Ờ': 'O', 'Ớ': 'O', 'Ở': 'O', 'Ỡ': 'O', 'Ợ': 'O',
        'Ù': 'U', 'Ú': 'U', 'Ủ': 'U', 'Ũ': 'U', 'Ụ': 'U',
        'Ư': 'U', 'Ừ': 'U', 'Ứ': 'U', 'Ử': 'U', 'Ữ': 'U', 'Ự': 'U',
        'Ỳ': 'Y', 'Ý': 'Y', 'Ỷ': 'Y', 'Ỹ': 'Y', 'Ỵ': 'Y'
        }
        for char, replacement in replacements.items():
            text = text.replace(char, replacement)
        return text

def normalize_text(text):
    """
    Chuẩn hóa chuỗi bằng cách chuyển thành không dấu, loại bỏ ký tự đặc biệt,
    và chuyển về chữ thường.
    """
    # Chuyển thành không dấu
    text = remove_vietnamese_diacritics(text)
    # Loại bỏ các ký tự không phải chữ cái hoặc chữ số
    normalized_text = re.sub(r'[^a-zA-Z0-9\s]', '', text)
    # Thay thế nhiều khoảng trắng bằng một khoảng trắng và loại bỏ khoảng trắng ở đầu/cuối
    normalized_text = re.sub(r'\s+', ' ', normalized_text).strip()
    return normalized_text.lower()
def get_form_code_with_retries(prompt_chon_dang, subject, level, content, images, last_key_index, max_retries=3):
    """
    Hàm riêng để gọi API lấy dạng câu hỏi và thử lại nếu thất bại.
    Trả về form_code và chỉ mục key cuối cùng được sử dụng.
    """
    form_code = None
    retries = 0
   
    while form_code is None and retries < max_retries:
            
            print( f"Thử lại lần {retries + 1} để lấy dạng câu hỏi...")
            time.sleep(2)
            api_response, last_key_index = call_gemini_api_with_images(
                prompt_chon_dang, subject, level, content, 
                 images, last_key_index
            )
            form_code = extract_form_code(api_response)
            print( f"Kết quả API (dạng câu hỏi) lần thử {retries + 1}: '{form_code}'")
            
            # Kiểm tra nếu lấy được form_code hợp lệ thì thoát vòng lặp
            if form_code and form_code != 'DD':
                break
            retries += 1
    
    # Trả về form_code và chỉ mục key cuối cùng được sử dụng
    return form_code, last_key_index

def process_word_file(input_path, subject, level, json_structure_path, temp_dir="temp_images"):
    global global_image_counter
    global_image_counter = 0

    print(f"[{get_timestamp()}] [START] Bắt đầu xử lý file: '{input_path}'")

   
    doc = docx.Document(input_path)
    print(f"[{get_timestamp()}] [INFO] Đã tải tài liệu đầu vào: '{input_path}'")
    
    output_filename = os.path.splitext(input_path)[0] + "_done.docx"
    if os.path.exists(output_filename):
        os.remove(output_filename)
        print(f"[{get_timestamp()}] [INFO] Đã xóa file đầu ra cũ: '{output_filename}'")
    output_doc = docx.Document()
    
    json_data = load_json_structure(json_structure_path)
    
    # Tạo một danh sách các câu hỏi với thông tin đầy đủ, bao gồm tiêu đề phần và chỉ mục
    all_questions_with_info = []
    for section in json_data['extracted_sections']:
        for i, question in enumerate(section['questions']):
            # Tạo một định danh duy nhất cho mỗi câu hỏi, ví dụ: "Tiêu đề câu hỏi - chỉ mục"
            unique_question_id = f"{question['question_block_start']}_{len(all_questions_with_info)}"
            all_questions_with_info.append({
                'unique_id': unique_question_id,
                'type_ques': question['type_ques'],
                'section_title': section['section_title'],
                'question_block_start': question['question_block_start']
            })
    
    total_questions = len(all_questions_with_info)
    processed_questions = 0
    questions_not_found = []
    last_used_key_index = None
    print(f"[{get_timestamp()}] [INFO] Đã tải cấu trúc JSON với tổng số {total_questions} câu hỏi.")

    prompt_chon_dang = read_prompt('prompt_chonDang.txt')
    prompt_dang_tn = read_prompt('prompt_dangTN.txt')
    prompt_dang_ds = read_prompt('prompt_dangDS.txt')
    prompt_dang_dien = read_prompt('prompt_dangDien.txt')
    prompt_dang_dd = read_prompt('prompt_dangDD.txt')

    # Sử dụng dictionary để ánh xạ ID duy nhất của câu hỏi đến chỉ mục đoạn văn
    paragraph_indices = {}
    print(f"\n[{get_timestamp()}] [INFO] Bắt đầu quét tài liệu để tìm chỉ mục đoạn văn...")
    
    # Lặp qua các đoạn văn bản
    current_question_index = 0
    for i, para in enumerate(doc.paragraphs):
        if current_question_index >= total_questions:
            break
        # print(f" [DEBUG] Quét đoạn văn {i}: '{para.text}'")
        normalized_para_text = normalize_text(para.text)
        current_question = all_questions_with_info[current_question_index]
        normalized_question_title = normalize_text(current_question['question_block_start'])
        
        # So sánh nội dung đoạn văn với tiêu đề câu hỏi tiếp theo trong danh sách
        if normalized_question_title in normalized_para_text:
            paragraph_indices[current_question['unique_id']] = i
            # print(f" [DEBUG] Đã tìm thấy chỉ mục {i} cho câu hỏi: '{normalized_question_title}'   trong doan văn {normalized_para_text}")
            current_question_index += 1
            
    print(f"\n[{get_timestamp()}] [INFO] Hoàn thành việc quét tài liệu.")
    
    
    processed_sections = set()
    
    # Lặp qua danh sách câu hỏi đã được tạo với ID duy nhất
    for i, question in enumerate(all_questions_with_info):
        question_title = question['question_block_start']
        section_title = question['section_title']
        unique_id = question['unique_id']
        question_number = i + 1
        form_code = question['type_ques']
        print(f"\n[{get_timestamp()}] [INFO] Đang xử lý câu hỏi số {question_number}/{total_questions}: {question_title} (ID: {unique_id})")
        
        start_index = paragraph_indices.get(unique_id)
        
        if start_index is None:
            print(f"[{get_timestamp()}] [WARNING] Không tìm thấy đoạn văn cho câu hỏi: '{question_title}'. Đánh dấu để phân tích.")
            questions_not_found.append(f"Câu {question_number}: {question_title}")
            continue
        
        if i + 1 < len(all_questions_with_info):
            next_question_id = all_questions_with_info[i+1]['unique_id']
            end_index = paragraph_indices.get(next_question_id)
            if end_index is None:
                end_index = len(doc.paragraphs)
                print(f"[{get_timestamp()}] [DEBUG] Không tìm thấy câu hỏi tiếp theo, đặt end_index = cuối tài liệu ({end_index}).")
            else:
                print(f"[{get_timestamp()}] [DEBUG] Tìm thấy chỉ mục kết thúc: {end_index} (câu hỏi tiếp theo).")
        else:
            end_index = len(doc.paragraphs)
            print(f"[{get_timestamp()}] [DEBUG] Đây là câu hỏi cuối cùng, đặt end_index = cuối tài liệu ({end_index}).")
        
        # Thêm tiêu đề phần nếu chưa được thêm
        if section_title and section_title not in processed_sections:
            output_doc.add_paragraph(section_title, 'Heading 1')
            processed_sections.add(section_title)
            
        # Thêm tiêu đề câu hỏi
        question_paragraph = output_doc.add_paragraph()
        question_paragraph.add_run(question_title).bold = True

        question_content_with_media, extracted_images = extract_question_content_with_media(
            doc, start_index, end_index, temp_dir
        )
        print(f"[{get_timestamp()}] [DEBUG] Trích xuất nội dung từ đoạn {start_index} đến {end_index}.")
        
        plain_text_content = "\n".join([doc.paragraphs[j].text for j in range(start_index, end_index)])
        print(plain_text_content)
        if has_valid_solution(plain_text_content):
            print(f"[{get_timestamp()}] [INFO] Đã có lời giải hợp lệ. Sao chép nội dung gốc.")
            copy_content_in_range(doc, output_doc, start_index, end_index)
        else:
            
            print(f"[{get_timestamp()}] [DEBUG] Kết quả API (dạng câu hỏi): '{form_code}'")
            
            if form_code == 'TN':
                prompt = prompt_dang_tn
            elif form_code == 'ĐS':
                prompt = prompt_dang_ds
            elif form_code == 'TLN':
                prompt = prompt_dang_dien
            elif form_code == 'DD':
                prompt = prompt_dang_dd
            else:
                form_code, last_used_key_index = get_form_code_with_retries(
                    prompt_chon_dang, subject, level, question_content_with_media, 
                     extracted_images, last_used_key_index, max_retries=3
                )
            
            print(f"[{get_timestamp()}] [INFO] Đã chọn prompt phù hợp. Gọi API lần 2.")
            
            final_response, last_used_key_index = call_gemini_api_with_images(
                prompt, subject, level, question_content_with_media, 
                 extracted_images, last_used_key_index
            )
            if form_code == 'ĐS' and final_response and '####' in final_response:
                parts = final_response.split('####', 1)
                question_part = parts[0]
                solution_part = parts[1]

                # Remove list markers from the solution part which is below '####'
                cleaned_solution = re.sub(r'^\s*[a-zA-Z][.)]\s*', '', solution_part, flags=re.MULTILINE)
                if not cleaned_solution.startswith('\n'):
                    cleaned_solution = '\n' + cleaned_solution
                final_response = question_part + '####' + cleaned_solution
            process_api_response_with_media(final_response, output_doc, extracted_images, images_folder=temp_dir)
        
        processed_questions += 1
        progress_percentage = (processed_questions / total_questions) * 100
        print(f"[{get_timestamp()}] [INFO] Tiến độ xử lý: {progress_percentage:.2f}% ({processed_questions}/{total_questions} câu)")
    
    if start_index is not None and end_index is not None and end_index < len(doc.paragraphs):
        print(f"[{get_timestamp()}] [INFO] Sao chép nội dung cuối tài liệu...")
        copy_content_in_range(doc, output_doc, end_index, len(doc.paragraphs))

    cleanup_temp_files(temp_dir) 
    
    try:
        output_doc.save(output_filename)
        print(f"[{get_timestamp()}] [SUCCESS] Đã lưu tài liệu hoàn chỉnh: '{output_filename}'")
        md_path = convert_docx_to_md(output_filename,temp_dir=temp_dir)
        md2docx(md_path)
        print("\n" + "="*50)
        print("[REPORT] Xử lý hoàn tất: 100%")
        print(f"Tổng số câu hỏi trong JSON: {total_questions}")
        print(f"Số câu hỏi đã xử lý thành công: {processed_questions}")
        print(f"Số câu hỏi không tìm thấy: {len(questions_not_found)}")
        if questions_not_found:
            print("\nDanh sách các câu hỏi không tìm thấy:")
            for q in questions_not_found:
                print(f"- {q}")
        print("="*50)
    
    except Exception as e:
        print(f"[{get_timestamp()}] [ERROR] Lỗi khi lưu tài liệu cuối cùng: {e}")
    return output_filename
def main():
    input_file = r"C:\Users\Admin\Downloads\chuyendangcauhoi\ChuyenDangCauHoi\doc\test_v2.docx"
    json_structure_file = r"C:\Users\Admin\Downloads\chuyendangcauhoi\ChuyenDangCauHoi\test_v2_split_parts\ai_struc_detec\part_1.json"
    subject = "Toán"
    level = "Lớp 8"
    api_keys_json_path = r"api_keys.json"
    
    try:
        process_word_file(input_file, subject, level, json_structure_file, temp_dir="Toán lớp 5 tập 1")
        print("Quy trình xử lý hoàn tất.")
    except Exception as e:
        print(f"Lỗi: {str(e)}")

if __name__ == "__main__":
    main()
