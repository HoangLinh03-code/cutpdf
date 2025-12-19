import os
import pypandoc
from typing import List
import re
from docxcompose.composer import Composer
def convert_md_to_docx(md_path: str):
    """
    Chuyển đổi một file .md sang .docx với các yêu cầu cụ thể.

    - Công thức toán LaTeX ($$...$$) được chuyển thành oMath trong .docx.
    - Hình ảnh trong thư mục 'media' được nhúng vào file .docx.
    - File output có cùng tên với file input.

    Args:
        md_path (str): Đường dẫn đến file .md cần chuyển đổi.
    """
    # Kiểm tra xem file Markdown có tồn tại không
    if not os.path.exists(md_path):
        print(f"Lỗi: File '{md_path}' không tồn tại.")
        return

    # Tách tên file và phần mở rộng để tạo tên file output
    base_name = os.path.splitext(md_path)[0]
    docx_path = base_name + 'convert.docx'

    print(f"Bắt đầu chuyển đổi file: '{md_path}'...")

    # Các tham số bổ sung cho Pandoc
    extra_args = [
        '--from=markdown',  # Đọc Markdown với công thức LaTeX
        '--to=docx',        # Đầu ra là định dạng docx
        # tham chiếu đến file docx đã có để giữ nguyên định dạng
    ]

    try:
        # Gọi hàm chuyển đổi của pypandoc
        pypandoc.convert_file(
            source_file=md_path,
            to='docx',
            outputfile=docx_path,
            extra_args=extra_args,
            # encoding='utf-8'  # Đảm bảo xử lý tốt tiếng Việt
        )
        print("-" * 30)
        print("🎉 Chuyển đổi thành công!")
        print(f"✔️ File Docx đã được lưu tại: '{docx_path}'")
        print("-" * 30)

    except (OSError, RuntimeError) as e:
        print("\n" + "*" * 50)
        print("LỖI: Có vẻ như Pandoc chưa được cài đặt hoặc không tìm thấy.")
        print("Vui lòng truy cập https://pandoc.org/installing.html để cài đặt Pandoc.")
        print(f"Chi tiết lỗi: {e}")
        print("*" * 50 + "\n")
    except Exception as e:
        print(f"Đã xảy ra lỗi không xác định trong quá trình chuyển đổi: {e}")

def clean_markdown_math_display(input_filepath):
    """
    Đọc một file Markdown, tìm và chuyển đổi các chuỗi dạng "$...$" thành "$ ... $"
    và xử lý các dấu thoát '\\' thừa bên trong các lệnh LaTeX như '\\times' thành '\times'.

    Args:
        input_filepath (str): Đường dẫn đến file Markdown đầu vào.
    """
    try:
        with open(input_filepath, 'r', encoding='utf-8') as infile:
            content = infile.read()

        def process_math_content(match):
            content_inside = match.group(1)
            content_inside = content_inside.strip()
            # content_inside = re.sub(r'\\(\\[a-zA-Z]+)', r'\1', content_inside)
            content_inside = re.sub(r'\\\\', r'\\', content_inside)
            return f"${content_inside}$"

        # Regex chính để tìm \$...\$
        final_cleaned_content = re.sub(r'\\\$([^\$]*?)\\\$', process_math_content, content)

        # Lưu nội dung đã được làm sạch vào file mới
        with open(input_filepath, 'w', encoding='utf-8') as outfile:
            outfile.write(final_cleaned_content)

        print(f"File đã được làm sạch và lưu tại: {input_filepath}")

    except FileNotFoundError:
        print(f"Lỗi: Không tìm thấy file tại đường dẫn {input_filepath}")
    except Exception as e:
        print(f"Đã xảy ra lỗi: {e}")

def md2docx(input_file: str):
    # Thay 'example.md' bằng tên file Markdown của bạn
    input_md_file = input_file
    
    # Gọi hàm làm sạch
    clean_markdown_math_display(input_md_file)
    
    if not os.path.exists(input_file):
        print(f"Không tìm thấy file '{input_file}'. Bạn hãy tạo một file .md với tên đó để thử nghiệm.")
    else:
        convert_md_to_docx(input_md_file)


import os
from docx import Document


def process_batch_multi_docx_to_single_docx_simple(input_folder: str = ''):
    """
    Hàm đơn giản để hợp nhất các file DOCX với hình ảnh sử dụng docxcompose.
    
    Cần cài đặt: pip install docxcompose
    
    Args:
        input_folder (str): Đường dẫn đến thư mục chứa các file .docx nguồn.
    """
    # Kiểm tra xem thư mục đầu vào có tồn tại không
    if not os.path.isdir(input_folder):
        print(f"Lỗi: Thư mục '{input_folder}' không tồn tại.")
        return

    # Lấy danh sách tất cả các file .docx trong thư mục
    all_docx_files = [
        os.path.join(input_folder, f) 
        for f in os.listdir(input_folder) 
        if f.endswith('.docx') and 'doneconvert' in f
    ]

    # Sắp xếp các file theo tên
    docx_files_to_merge = sorted(all_docx_files, key=os.path.basename)

    # Kiểm tra nếu không tìm thấy file DOCX nào
    if not docx_files_to_merge:
        print(f"Không tìm thấy file DOCX nào trong thư mục '{input_folder}'.")
        return

    # Xác định tên file DOCX đầu ra
    output_docx_name = os.path.join(input_folder, f"{os.path.basename(input_folder)}_merged.docx")

    print(f"Tìm thấy các file DOCX để hợp nhất: {docx_files_to_merge}")

    try:
        # Sử dụng file đầu tiên làm base document
        master_doc = Document(docx_files_to_merge[0])
        composer = Composer(master_doc)
        
        print("Bắt đầu hợp nhất các file DOCX...")
        
        # Thêm các file còn lại vào master document
        for file_path in docx_files_to_merge[1:]:
            if os.path.exists(file_path):
                try:
                    doc_to_append = Document(file_path)
                    composer.append(doc_to_append)
                    print(f"Đã hợp nhất nội dung từ file: {os.path.basename(file_path)}")
                except Exception as e:
                    print(f"Cảnh báo: Không thể hợp nhất file '{os.path.basename(file_path)}'. Lỗi: {e}")
            else:
                print(f"Cảnh báo: File '{os.path.basename(file_path)}' không tồn tại.")

        # Lưu document đã được hợp nhất
        composer.save(output_docx_name)
        
        print("-" * 30)
        print("🎉 Hợp nhất thành công!")
        print(f"✔️ File DOCX đã được lưu tại: '{output_docx_name}'")
        print("-" * 30)

    except ImportError:
        print("⚠️ Cần cài đặt thư viện docxcompose:")
        print("   pip install docxcompose")
        print()
        print("Đang sử dụng phương pháp thay thế...")
        
        # Fallback method - sử dụng cách thủ công
        process_batch_multi_docx_to_single_docx_manual(input_folder)
        
    except Exception as e:
        print(f"Đã xảy ra lỗi không xác định trong quá trình hợp nhất: {e}")


def process_batch_multi_docx_to_single_docx_manual(input_folder: str = ''):
    """
    Phương pháp thủ công để hợp nhất DOCX - copy từng paragraph và run
    """
    # Kiểm tra thư mục
    if not os.path.isdir(input_folder):
        print(f"Lỗi: Thư mục '{input_folder}' không tồn tại.")
        return

    # Lấy danh sách file
    all_docx_files = [
        os.path.join(input_folder, f) 
        for f in os.listdir(input_folder) 
        if f.endswith('.docx') and 'doneconvert' in f
    ]

    docx_files_to_merge = sorted(all_docx_files, key=os.path.basename)

    if not docx_files_to_merge:
        print(f"Không tìm thấy file DOCX nào trong thư mục '{input_folder}'.")
        return

    output_docx_name = os.path.join(input_folder, f"{os.path.basename(input_folder)}_merged.docx")

    try:
        # Tạo document mới
        merged_document = Document()

        print("Bắt đầu hợp nhất các file DOCX (phương pháp thủ công)...")
        
        for i, file_path in enumerate(docx_files_to_merge):
            if os.path.exists(file_path):
                try:
                    source_document = Document(file_path)
                    
                    # Thêm page break giữa các document (trừ document đầu tiên)
                    if i > 0:
                        merged_document.add_page_break()
                    
                    # Copy từng paragraph
                    for paragraph in source_document.paragraphs:
                        # Tạo paragraph mới
                        new_paragraph = merged_document.add_paragraph()
                        new_paragraph.alignment = paragraph.alignment
                        
                        # Copy từng run để giữ nguyên định dạng
                        for run in paragraph.runs:
                            new_run = new_paragraph.add_run(run.text)
                            
                            # Copy định dạng
                            if run.bold:
                                new_run.bold = True
                            if run.italic:
                                new_run.italic = True
                            if run.underline:
                                new_run.underline = True
                            if run.font.size:
                                new_run.font.size = run.font.size
                            if run.font.name:
                                new_run.font.name = run.font.name
                    
                    # Copy tables
                    for table in source_document.tables:
                        # Tạo table mới với cùng số hàng và cột
                        new_table = merged_document.add_table(
                            rows=len(table.rows), 
                            cols=len(table.columns)
                        )
                        
                        # Copy nội dung từng cell
                        for row_idx, row in enumerate(table.rows):
                            for col_idx, cell in enumerate(row.cells):
                                new_table.cell(row_idx, col_idx).text = cell.text
                    
                    print(f"Đã hợp nhất nội dung từ file: {os.path.basename(file_path)}")
                    
                except Exception as e:
                    print(f"Cảnh báo: Không thể hợp nhất file '{os.path.basename(file_path)}'. Lỗi: {e}")
            else:
                print(f"Cảnh báo: File '{os.path.basename(file_path)}' không tồn tại.")

        # Lưu file
        merged_document.save(output_docx_name)
        
        print("-" * 30)
        print("🎉 Hợp nhất thành công!")
        print(f"✔️ File DOCX đã được lưu tại: '{output_docx_name}'")
        print("ℹ️ Lưu ý: Hình ảnh có thể không được copy trong phương pháp này.")
        print("   Để copy hình ảnh, cài đặt: pip install docxcompose")
        print("-" * 30)

    except Exception as e:
        print(f"Đã xảy ra lỗi không xác định trong quá trình hợp nhất: {e}")


# Function chính - thử docxcompose trước, fallback sang manual
def process_batch_multi_docx_to_single_docx(input_folder: str = ''):
    """
    Hàm chính để hợp nhất DOCX với khả năng copy hình ảnh
    """
    try:
        # Thử sử dụng docxcompose trước
        process_batch_multi_docx_to_single_docx_simple(input_folder)
    except ImportError:
        # Nếu không có docxcompose, sử dụng phương pháp thủ công
        print("Sử dụng phương pháp thủ công...")
        process_batch_multi_docx_to_single_docx_manual(input_folder)
# md2docx(r"C:\Users\Admin\Downloads\Xử lý chuyển xml (3)\Xử lý chuyển xml (3)\a.md")
# Hàm không còn được sử dụng trong logic mới, nhưng vẫn được giữ lại để tương thích.
# process_batch_multi_md_to_single_docx(r"D:\Xử lý chuyển xml\Giải vở bài tập tiếng việt lớp 5 - VBT Tiếng Việt_split_parts")

# process_batch_multi_docx_to_single_docx(r"C:\Users\Admin\Downloads\Xử lý chuyển xml\Toán lớp 5 tập 2 - Kết nối tri thức_parts_docx_files")